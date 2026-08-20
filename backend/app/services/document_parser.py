import re
import unicodedata
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

PARSER_VERSION = "fitcv-parser-v5-gemini-source"
MAX_CV_BYTES = 10 * 1024 * 1024

SKILL_ALIASES: dict[str, tuple[str, ...]] = {
    "Python": ("python",),
    "Java": ("java",),
    "JavaScript": ("javascript", "js"),
    "TypeScript": ("typescript", "ts"),
    "C++": ("c++",),
    "C#": ("c#", "c sharp"),
    "Go": ("golang",),
    "PHP": ("php",),
    "HTML": ("html", "html5"),
    "CSS": ("css", "css3"),
    "React": ("react.js", "reactjs", "react"),
    "Node.js": ("node.js", "nodejs", "node js"),
    "ASP.NET": ("asp.net", "asp net", "aspnet"),
    "FastAPI": ("fastapi",),
    "Django": ("django",),
    "Flask": ("flask",),
    "Laravel": ("laravel",),
    "SQL": ("sql",),
    "MySQL": ("mysql",),
    "SQL Server": ("sql server",),
    "PostgreSQL": ("postgresql", "postgres"),
    "MongoDB": ("mongodb", "mongo db"),
    "Redis": ("redis",),
    "REST APIs": ("rest api", "restful api", "restful services"),
    "GraphQL": ("graphql",),
    "Microservices": ("microservices", "microservice architecture"),
    "Docker": ("docker",),
    "Kubernetes": ("kubernetes", "k8s"),
    "AWS": ("amazon web services", "aws"),
    "Azure": ("microsoft azure", "azure"),
    "Google Cloud": ("google cloud platform", "google cloud", "gcp"),
    "Git": ("git",),
    "CI/CD": ("ci/cd", "continuous integration", "continuous delivery"),
    "Linux": ("linux",),
    "Bash": ("bash", "shell scripting"),
    "Tailwind CSS": ("tailwind css",),
    "Bootstrap": ("bootstrap",),
    "Figma": ("figma",),
    "Postman": ("postman",),
    "XAMPP": ("xampp",),
    "JWT": ("jwt",),
    "VNPay": ("vnpay",),
    "Agile": ("agile",),
    "Scrum": ("scrum",),
    "Machine Learning": ("machine learning",),
    "Active Directory": ("active directory",),
    "Burp Suite": ("burp suite",),
    "Cybersecurity": ("cybersecurity", "cyber security"),
    "Digital Forensics": ("digital forensics",),
    "Firewalls": ("firewall", "firewalls"),
    "Incident Response": ("incident response",),
    "IDS/IPS": ("ids/ips", "intrusion detection", "intrusion prevention"),
    "MITRE ATT&CK": ("mitre att&ck", "mitre attack"),
    "Microsoft Sentinel": ("microsoft sentinel", "azure sentinel"),
    "Network Security": ("network security",),
    "Nessus": ("nessus",),
    "Nmap": ("nmap",),
    "OWASP Top 10": ("owasp top 10", "owasp top ten"),
    "SIEM": ("siem", "security information and event management"),
    "SOC": ("security operations center", "soc"),
    "Splunk": ("splunk",),
    "TCP/IP": ("tcp/ip", "tcp ip"),
    "Vulnerability Assessment": ("vulnerability assessment", "vulnerability scanning"),
    "Wireshark": ("wireshark",),
}

SOFT_SKILLS: dict[str, tuple[str, ...]] = {
    "Communication": (
        "communication",
        "communicate",
        "verbal and written communication",
        "interpersonal skills",
        "giao tiếp",
        "kỹ năng giao tiếp",
        "truyền đạt",
    ),
    "Teamwork": (
        "teamwork",
        "team player",
        "team-oriented",
        "collaboration",
        "collaborative",
        "cross-functional collaboration",
        "làm việc nhóm",
        "kỹ năng làm việc nhóm",
        "hợp tác",
        "phối hợp",
    ),
    "Leadership": (
        "leadership",
        "lead a team",
        "mentoring",
        "lãnh đạo",
        "dẫn dắt",
        "quản lý nhóm",
    ),
    "Problem Solving": (
        "problem solving",
        "problem-solving",
        "analytical thinking",
        "analytical skills",
        "giải quyết vấn đề",
        "xử lý vấn đề",
        "tư duy phân tích",
    ),
    "Adaptability": (
        "adaptability",
        "adaptable",
        "flexibility",
        "thích nghi",
        "linh hoạt",
    ),
    "Time Management": (
        "time management",
        "prioritization",
        "quản lý thời gian",
        "quan ly thoi gian",
        "sắp xếp công việc",
    ),
    "Critical Thinking": (
        "critical thinking",
        "tư duy phản biện",
        "tư duy logic",
    ),
    "Creativity": (
        "creativity",
        "creative thinking",
        "sáng tạo",
        "sang tao",
        "tư duy sáng tạo",
    ),
    "Attention to Detail": (
        "attention to detail",
        "detail-oriented",
        "detail oriented",
        "tỉ mỉ",
        "ti mi",
        "cẩn thận",
        "can than",
    ),
    "Self-Learning": (
        "self-learning",
        "self learning",
        "self-taught",
        "fast learner",
        "quick learner",
        "willing to learn",
        "eager to learn",
        "tự học",
        "tu hoc",
        "tự nghiên cứu",
        "cầu tiến",
        "cau tien",
        "ham học hỏi",
    ),
    "Responsibility": (
        "sense of responsibility",
        "responsibility",
        "ownership",
        "accountable",
        "trách nhiệm",
        "trach nhiem",
        "có trách nhiệm",
        "tinh thần trách nhiệm",
    ),
    "Work Under Pressure": (
        "work under pressure",
        "working under pressure",
        "work well under pressure",
        "chịu được áp lực",
        "chiu duoc ap luc",
        "chịu áp lực",
        "chiu ap luc",
        "làm việc dưới áp lực",
        "áp lực công việc",
    ),
    "Presentation": (
        "presentation skills",
        "public speaking",
        "thuyết trình",
        "thuyet trinh",
        "kỹ năng thuyết trình",
    ),
    "Proactiveness": (
        "proactive",
        "self-motivated",
        "self driven",
        "self-driven",
        "chủ động",
        "chu dong",
        "năng động",
        "nang dong",
        "nhiệt tình",
    ),
    "Negotiation": (
        "negotiation",
        "đàm phán",
        "dam phan",
        "thương lượng",
    ),
    "Organization": (
        "organizational skills",
        "well-organized",
        "well organized",
        "tổ chức công việc",
        "khoa học",
        "ngăn nắp",
    ),
}


def validate_cv_content(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        if not content.startswith(b"%PDF-"):
            raise ValueError("The uploaded file is not a valid PDF.")
        return "PDF"
    if suffix == ".docx":
        try:
            with ZipFile(BytesIO(content)) as archive:
                names = set(archive.namelist())
        except BadZipFile as exc:
            raise ValueError("The uploaded file is not a valid DOCX document.") from exc
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ValueError("The uploaded file is not a valid DOCX document.")
        return "DOCX"
    raise ValueError("Only PDF and DOCX files are supported.")


def extract_document_text(file_path: Path, file_type: str) -> str:
    if not file_path.is_file():
        raise FileNotFoundError(
            "The uploaded CV file is no longer available on the server. "
            "Please upload the CV again."
        )
    if file_type == "PDF":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF parsing is unavailable; install backend requirements.") from exc
        text = "\n".join(page.extract_text() or "" for page in PdfReader(str(file_path)).pages)
        normalized = preprocess_document_text(text)
        if len(normalized) < 20:
            from app.services.ocr_service import OcrError, extract_pdf_text

            try:
                normalized = preprocess_document_text(extract_pdf_text(file_path))
            except OcrError as exc:
                raise ValueError(str(exc)) from exc
    elif file_type == "DOCX":
        from docx import Document

        document = Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_rows = [" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows]
        text = "\n".join([*paragraphs, *table_rows])
        normalized = preprocess_document_text(text)
    else:
        raise ValueError("Unsupported CV file type.")

    if len(normalized) < 20:
        raise ValueError("No readable CV text was found after OCR.")
    return normalized


def parse_cv_text(text: str) -> dict:
    normalized = preprocess_document_text(text)
    if len(normalized) < 20:
        raise ValueError("CV text is empty or too short to parse.")
    return {
        "skills": _extract_terms(normalized, SKILL_ALIASES),
        "experience_years": _extract_years(normalized),
        "education": _extract_education(normalized),
        "soft_skills": _extract_terms(normalized, SOFT_SKILLS),
        "sections": _extract_sections(normalized),
    }


def parse_jd_text(text: str) -> dict:
    normalized = preprocess_document_text(text)
    if len(normalized) < 50:
        raise ValueError("Job description must contain at least 50 readable characters.")

    required: set[str] = set()
    preferred: set[str] = set()
    chunks = re.split(r"\n+|(?<=[.!?])\s+", normalized)
    for chunk in chunks:
        found = set(_extract_terms(chunk, SKILL_ALIASES))
        if not found:
            continue
        if re.search(r"\b(preferred|nice to have|bonus|plus|advantage)\b", chunk, re.IGNORECASE):
            preferred.update(found)
        else:
            required.update(found)

    preferred.difference_update(required)
    return {
        "required_skills": sorted(required),
        "preferred_skills": sorted(preferred),
        "experience_years": _extract_years(normalized),
        "education": _extract_education(normalized),
        "soft_skills": _extract_terms(normalized, SOFT_SKILLS),
    }


def preprocess_document_text(value: str) -> str:
    prepared = unicodedata.normalize(
        "NFKC",
        value.replace("\x00", "").replace("\u00ad", ""),
    )
    prepared = re.sub(r"(?<=[A-Za-z])-\s*\n\s*(?=[a-z])", "", prepared)
    return _normalize_text(prepared)


def _normalize_text(value: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _extract_terms(text: str, aliases: dict[str, tuple[str, ...]]) -> list[str]:
    lowered = text.casefold()
    found = {
        canonical
        for canonical, values in aliases.items()
        if any(re.search(rf"(?<!\w){re.escape(alias.casefold())}(?!\w)", lowered) for alias in values)
    }
    return sorted(found)


def _extract_years(text: str) -> float | None:
    values = [
        float(value)
        for value in re.findall(r"\b(\d{1,2}(?:\.\d)?)\s*\+?\s*(?:years?|yrs?)\b", text, re.IGNORECASE)
        if float(value) <= 50
    ]
    return max(values) if values else None


def _extract_education(text: str) -> str | None:
    lowered = text.casefold()
    levels = (
        (
            "Doctorate",
            ("phd", "ph.d", "doctorate", "doctor of philosophy"),
        ),
        (
            "Master",
            (
                "master's",
                "masters degree",
                "master degree",
                "master of",
                "master in",
                "msc",
                "m.sc",
                "mba",
            ),
        ),
        (
            "Bachelor",
            (
                "bachelor's",
                "bachelors degree",
                "bachelor degree",
                "bachelor of",
                "bachelor in",
                "bsc",
                "b.sc",
                "beng",
                "b.eng",
            ),
        ),
        ("Associate", ("associate degree",)),
        ("High School", ("high school",)),
    )
    return next((level for level, aliases in levels if any(alias in lowered for alias in aliases)), None)


def _extract_sections(text: str) -> dict[str, str]:
    headers = {
        "summary": "summary",
        "professional summary": "summary",
        "experience": "experience",
        "professional experience": "experience",
        "work experience": "experience",
        "employment history": "experience",
        "education": "education",
        "skills": "skills",
        "technical skills": "skills",
        "core competencies": "skills",
        "projects": "projects",
        "selected projects": "projects",
    }
    sections: dict[str, list[str]] = {}
    current = "other"
    for line in text.splitlines():
        candidate = line.rstrip(":").strip().casefold()
        if candidate in headers:
            current = headers[candidate]
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)
    return {name: "\n".join(lines).strip() for name, lines in sections.items() if lines}
