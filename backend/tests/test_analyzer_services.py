import json
import unittest
import base64
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import MagicMock, patch

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.deps import get_current_account
from app.core.config import settings
from app.db.session import Base, get_db
from app.main import app
from app.models.account import Account, AuthProvider
from app.models.analyzer import Cv, CvParseResult, JdParseResult, Job, JobDescription, MatchResult
from app.models.improvement import AiTask
from app.repositories import ai_tasks, analyzer
from app.services.document_parser import (
    PARSER_VERSION,
    extract_document_text,
    parse_cv_text,
    parse_jd_text,
    preprocess_document_text,
    validate_cv_content,
)
from app.services.matching_service import (
    ALGORITHM_VERSION,
    match_documents,
    supplement_semantic_cv,
)
from app.services.analyzer_service import _selected_analyzer_config, run_cv_parse
from app.services import ocr_service
from app.services.gemini_analyzer import (
    GeminiAnalyzerError,
    extract_cv_inputs_from_file,
    extract_cv_search_profile,
    extract_match_inputs,
)


class DocumentParserTests(unittest.TestCase):
    def test_extracts_build_cv_aliases_and_sections(self) -> None:
        cv = parse_cv_text(
            """CORE COMPETENCIES
            Backend Development using ASP.NET (C#) and PHP.
            EDUCATION
            Major: Information Technology - Saigon Technology University.
            PROFESSIONAL EXPERIENCE
            Internship project using ReactJS, HTML5, CSS3, REST APIs and Figma.
            SELECTED PROJECTS
            Built a payment flow with SQL Server, Tailwind CSS, Bootstrap, JWT,
            Postman, XAMPP and VNPay.
            """
        )

        self.assertIn("ASP.NET", cv["skills"])
        self.assertIn("HTML", cv["skills"])
        self.assertIn("CSS", cv["skills"])
        self.assertIn("SQL Server", cv["skills"])
        self.assertIn("Tailwind CSS", cv["skills"])
        self.assertIn("Bootstrap", cv["skills"])
        self.assertIn("Figma", cv["skills"])
        self.assertIn("Postman", cv["skills"])
        self.assertIn("JWT", cv["skills"])
        self.assertIn("VNPay", cv["skills"])
        self.assertIn("skills", cv["sections"])
        self.assertIn("experience", cv["sections"])
        self.assertIn("projects", cv["sections"])

    def test_extracts_shared_cv_and_jd_contract(self) -> None:
        cv = parse_cv_text(
            """Technical Skills
            Python, FastAPI, MySQL, Docker, Git
            Experience
            4 years building REST APIs in Agile teams.
            Education
            Bachelor's degree in Computer Science.
            Communication and teamwork."""
        )
        jd = parse_jd_text(
            """Backend Developer requirements: 3 years of experience with Python, FastAPI, MySQL and REST APIs.
            Bachelor's degree required. Strong communication and teamwork.
            Docker and Redis are nice to have for this position."""
        )

        self.assertIn("FastAPI", cv["skills"])
        self.assertIn("Redis", jd["preferred_skills"])
        self.assertNotIn("Redis", jd["required_skills"])
        self.assertEqual(jd["experience_years"], 3.0)

    def test_preprocesses_ocr_text_and_recognizes_master_of_science(self) -> None:
        text = preprocess_document_text(
            "EDUCATION\nMaster of Science in Artificial\n"
            "Intelligence\nProficient in profes-\nsional Python development."
        )
        parsed = parse_cv_text(text)

        self.assertIn("professional Python", text)
        self.assertEqual(parsed["education"], "Master")
        self.assertIn("Python", parsed["skills"])

    def test_rejects_fake_pdf(self) -> None:
        with self.assertRaisesRegex(ValueError, "valid PDF"):
            validate_cv_content("resume.pdf", b"not a pdf")

    def test_extracts_docx_text(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.docx"
            document = Document()
            document.add_heading("Technical Skills")
            document.add_paragraph("Python, FastAPI, MySQL")
            document.save(path)
            self.assertIn("FastAPI", extract_document_text(path, "DOCX"))

    def test_scanned_pdf_uses_ocr_fallback(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "scanned.pdf"
            path.write_bytes(b"%PDF-1.4\nscanned")
            page = MagicMock()
            page.extract_text.return_value = ""
            reader = MagicMock()
            reader.pages = [page]
            with (
                patch("pypdf.PdfReader", return_value=reader),
                patch(
                    "app.services.ocr_service.extract_pdf_text",
                    return_value=(
                        "Technical Skills\nPython FastAPI SQL\n"
                        "Experience\nThree years building APIs"
                    ),
                ) as ocr,
            ):
                text = extract_document_text(path, "PDF")

        self.assertIn("FastAPI", text)
        ocr.assert_called_once_with(path)


class OcrServiceTests(unittest.TestCase):
    def setUp(self) -> None:
        self.original_provider = settings.ocr_provider
        self.original_model = settings.ocr_model
        self.original_key = settings.gemini_api_key
        self.original_timeout = settings.ocr_timeout_seconds
        self.original_thinking_level = settings.gemini_thinking_level
        settings.ocr_provider = "gemini"
        settings.ocr_model = "gemini-ocr-test"
        settings.gemini_api_key = "test-key"
        settings.ocr_timeout_seconds = 7
        settings.gemini_thinking_level = "high"

    def tearDown(self) -> None:
        settings.ocr_provider = self.original_provider
        settings.ocr_model = self.original_model
        settings.gemini_api_key = self.original_key
        settings.ocr_timeout_seconds = self.original_timeout
        settings.gemini_thinking_level = self.original_thinking_level

    def test_sends_pdf_inline_and_returns_transcription(self) -> None:
        response = MagicMock()
        response.status_code = 200
        response.json.return_value = {
            "candidates": [
                {
                    "finishReason": "STOP",
                    "content": {
                        "parts": [
                            {
                                "text": (
                                    "Technical Skills\nPython, FastAPI\n"
                                    "Experience\n3 years"
                                )
                            }
                        ]
                    },
                }
            ]
        }
        with TemporaryDirectory() as directory:
            path = Path(directory) / "scan.pdf"
            pdf_bytes = b"%PDF-1.4\nimage-only"
            path.write_bytes(pdf_bytes)
            with patch(
                "app.services.ocr_service.requests.post",
                return_value=response,
            ) as post:
                text = ocr_service.extract_pdf_text(path)

        self.assertIn("FastAPI", text)
        request = post.call_args
        self.assertEqual(
            request.args[0],
            (
                "https://generativelanguage.googleapis.com/v1beta/models/"
                "gemini-ocr-test:generateContent"
            ),
        )
        self.assertEqual(request.kwargs["timeout"], 7)
        inline = request.kwargs["json"]["contents"][0]["parts"][0]["inlineData"]
        self.assertEqual(inline["mimeType"], "application/pdf")
        self.assertEqual(base64.b64decode(inline["data"]), pdf_bytes)
        self.assertEqual(
            request.kwargs["json"]["generationConfig"]["thinkingConfig"],
            {"thinkingLevel": "high"},
        )
        self.assertNotIn("temperature", request.kwargs["json"]["generationConfig"])

    def test_reports_ocr_finish_reason(self) -> None:
        with self.assertRaisesRegex(ocr_service.OcrError, "MAX_TOKENS"):
            ocr_service._output_text(
                {"candidates": [{"finishReason": "MAX_TOKENS", "content": {}}]}
            )

    def test_retries_transient_ocr_connection_error(self) -> None:
        response = MagicMock(status_code=200)
        response.json.return_value = {
            "candidates": [
                {
                    "finishReason": "STOP",
                    "content": {"parts": [{"text": "Python Engineer"}]},
                }
            ]
        }
        with (
            patch.object(settings, "gemini_max_retries", 1),
            patch(
                "app.services.ocr_service.requests.post",
                side_effect=[
                    ocr_service.requests.ConnectionError("temporary reset"),
                    response,
                ],
            ) as post,
            patch("app.services.ocr_service.time.sleep") as retry_sleep,
        ):
            payload = ocr_service._send_request(
                "https://example.test",
                {"contents": []},
            )

        self.assertEqual(post.call_count, 2)
        retry_sleep.assert_called_once_with(0.5)
        self.assertEqual(
            payload["candidates"][0]["content"]["parts"][0]["text"],
            "Python Engineer",
        )

    def test_requires_api_key_for_scanned_pdf_ocr(self) -> None:
        settings.gemini_api_key = None
        with TemporaryDirectory() as directory:
            path = Path(directory) / "scan.pdf"
            path.write_bytes(b"%PDF-1.4\nimage-only")
            with self.assertRaisesRegex(ocr_service.OcrError, "GEMINI_API_KEY"):
                ocr_service.extract_pdf_text(path)


class MatchingServiceTests(unittest.TestCase):
    def test_scores_evidence_and_probability(self) -> None:
        result = match_documents(
            {
                "skills": ["Python", "FastAPI", "MySQL", "REST APIs", "Docker"],
                "experience_years": 4.0,
                "education": "Bachelor",
                "soft_skills": ["Communication", "Teamwork"],
            },
            {
                "required_skills": ["Python", "FastAPI", "MySQL", "REST APIs"],
                "preferred_skills": ["Docker", "Redis"],
                "experience_years": 3.0,
                "education": "Bachelor",
                "soft_skills": ["Communication", "Teamwork"],
            },
        )

        self.assertEqual(result["match_label"], "Strong Match")
        self.assertGreaterEqual(result["overall_score"], 80)
        self.assertLessEqual(result["pass_probability"], 95)
        self.assertEqual(result["breakdown"]["skills"]["missing"], ["Redis"])

    def test_redistributes_weights_when_jd_omits_categories(self) -> None:
        result = match_documents(
            {"skills": [], "experience_years": None, "education": "Master", "soft_skills": []},
            {"required_skills": [], "preferred_skills": [], "experience_years": None, "education": "Bachelor", "soft_skills": []},
        )
        self.assertEqual(result["overall_score"], 100.0)

    def test_custom_job_weights_change_the_candidate_score(self) -> None:
        cv = {
            "skills": ["Python"],
            "experience_years": 1,
            "education": None,
            "soft_skills": [],
        }
        jd = {
            "required_skills": ["Python"],
            "preferred_skills": [],
            "experience_years": 4,
            "education": None,
            "soft_skills": [],
        }

        skills_first = match_documents(
            cv,
            jd,
            weights={
                "skills": 80,
                "experience": 20,
                "education": 0,
                "soft_skills": 0,
            },
        )
        experience_first = match_documents(
            cv,
            jd,
            weights={
                "skills": 20,
                "experience": 80,
                "education": 0,
                "soft_skills": 0,
            },
        )

        self.assertEqual(skills_first["overall_score"], 85.0)
        self.assertEqual(experience_first["overall_score"], 40.0)
        self.assertEqual(skills_first["scoring_weights"]["skills"], 80.0)

    def test_custom_weights_still_redistribute_missing_categories(self) -> None:
        result = match_documents(
            {"skills": ["Python"], "soft_skills": []},
            {
                "required_skills": ["Python"],
                "preferred_skills": [],
                "experience_years": None,
                "education": None,
                "soft_skills": [],
            },
            weights={
                "skills": 25,
                "experience": 50,
                "education": 15,
                "soft_skills": 10,
            },
        )

        self.assertEqual(result["overall_score"], 100.0)

    def test_rejects_invalid_custom_weights(self) -> None:
        with self.assertRaisesRegex(ValueError, "total 100"):
            match_documents(
                {"skills": ["Python"]},
                {"required_skills": ["Python"]},
                weights={
                    "skills": 50,
                    "experience": 30,
                    "education": 15,
                    "soft_skills": 10,
                },
            )

    def test_rejects_unscorable_jd(self) -> None:
        with self.assertRaisesRegex(ValueError, "no scorable"):
            match_documents(
                {"skills": [], "experience_years": None, "education": None, "soft_skills": []},
                {"required_skills": [], "preferred_skills": [], "experience_years": None, "education": None, "soft_skills": []},
            )

    def test_matches_canonical_skill_variants(self) -> None:
        result = match_documents(
            {
                "skills": ["REST APIs"],
                "experience_years": None,
                "education": None,
                "soft_skills": [],
            },
            {
                "required_skills": ["REST API"],
                "preferred_skills": [],
                "experience_years": None,
                "education": None,
                "soft_skills": [],
            },
        )
        self.assertEqual(result["breakdown"]["skills"]["score"], 100.0)
        self.assertEqual(result["breakdown"]["skills"]["matched"], ["REST API"])

    def test_one_of_group_is_satisfied_by_one_matching_skill(self) -> None:
        result = match_documents(
            {
                "skills": ["C++"],
                "experience_years": None,
                "education": None,
                "soft_skills": [],
            },
            {
                "required_skills": [],
                "preferred_skills": [],
                "required_skill_groups": [
                    {
                        "skills": ["C++", "Python", "C#", "Java"],
                        "minimum_required": 1,
                    }
                ],
                "experience_years": None,
                "education": None,
                "soft_skills": [],
            },
        )

        skills = result["breakdown"]["skills"]
        self.assertEqual(skills["score"], 100.0)
        self.assertEqual(skills["matched"], ["C++"])
        self.assertEqual(skills["missing"], [])
        self.assertTrue(skills["groups"][0]["satisfied"])

    def test_minimum_skill_group_reports_one_group_gap(self) -> None:
        result = match_documents(
            {"skills": ["Python"], "soft_skills": []},
            {
                "required_skills": [],
                "preferred_skills": [],
                "required_skill_groups": [
                    {
                        "skills": ["Python", "Java", "Go"],
                        "minimum_required": 2,
                    }
                ],
                "soft_skills": [],
            },
        )

        skills = result["breakdown"]["skills"]
        self.assertEqual(skills["score"], 50.0)
        self.assertEqual(
            skills["missing"],
            ["At least 2 of: Go, Java, Python"],
        )
        self.assertFalse(skills["groups"][0]["satisfied"])

    def test_supplements_semantic_cv_with_locally_parsed_terms(self) -> None:
        semantic_cv = {
            "skills": ["TensorFlow"],
            "experience_years": None,
            "education": None,
            "soft_skills": [],
        }
        parsed_cv = {
            "skills": ["Machine Learning", "Python"],
            "experience_years": None,
            "education": "Master",
            "soft_skills": ["Problem Solving"],
        }
        jd = {
            "required_skills": ["Python", "FastAPI"],
            "preferred_skills": [],
            "experience_years": None,
            "education": None,
            "soft_skills": [],
        }

        supplemented = supplement_semantic_cv(semantic_cv, parsed_cv)
        result = match_documents(supplemented, jd)

        self.assertIn("Python", supplemented["skills"])
        self.assertEqual(supplemented["education"], "Master")
        self.assertEqual(result["breakdown"]["skills"]["matched"], ["Python"])
        self.assertEqual(result["breakdown"]["skills"]["score"], 50.0)


class GeminiAnalyzerTests(unittest.TestCase):
    def setUp(self) -> None:
        self.original_provider = settings.analyzer_provider
        self.original_key = settings.gemini_api_key
        self.original_model = settings.gemini_model
        self.original_timeout = settings.gemini_timeout_seconds
        self.original_retries = settings.gemini_max_retries
        self.original_thinking_level = settings.gemini_thinking_level
        self.original_structured_thinking_level = settings.gemini_structured_thinking_level
        self.original_structured_output_tokens = settings.gemini_structured_output_tokens
        settings.analyzer_provider = "gemini"
        settings.gemini_api_key = "test-key"
        settings.gemini_model = "gemini-3.6-flash"
        settings.gemini_thinking_level = "high"
        settings.gemini_structured_thinking_level = "low"
        settings.gemini_structured_output_tokens = 24_000
        settings.gemini_timeout_seconds = 1
        settings.gemini_max_retries = 1

    def tearDown(self) -> None:
        settings.analyzer_provider = self.original_provider
        settings.gemini_api_key = self.original_key
        settings.gemini_model = self.original_model
        settings.gemini_timeout_seconds = self.original_timeout
        settings.gemini_max_retries = self.original_retries
        settings.gemini_thinking_level = self.original_thinking_level
        settings.gemini_structured_thinking_level = self.original_structured_thinking_level
        settings.gemini_structured_output_tokens = self.original_structured_output_tokens

    @staticmethod
    def _gemini_response(output: dict) -> MagicMock:
        response = MagicMock(status_code=200)
        response.json.return_value = {
            "candidates": [
                {
                    "content": {"parts": [{"text": json.dumps(output)}]},
                    "finishReason": "STOP",
                }
            ]
        }
        return response

    @staticmethod
    def _cv_file_output(
        *,
        education_entries: list[dict] | None = None,
        education: str | None = None,
        education_evidence: str | None = None,
    ) -> dict:
        return {
            "skills": [],
            "experience_years": None,
            "experience_evidence": None,
            "education": education,
            "education_evidence": education_evidence,
            "education_entries": education_entries or [],
            "experience_entries": [],
            "soft_skills": [],
        }

    @staticmethod
    def _match_output(
        *,
        cv_skills: list[dict] | None = None,
        jd_required_skills: list[dict] | None = None,
        jd_required_skill_groups: list[dict] | None = None,
    ) -> dict:
        return {
            "cv": {
                "skills": cv_skills or [],
                "experience_years": None,
                "experience_evidence": None,
                "education": None,
                "education_evidence": None,
                "education_entries": [],
                "experience_entries": [],
                "soft_skills": [],
            },
            "jd": {
                "required_skills": jd_required_skills or [],
                "preferred_skills": [],
                "required_skill_groups": jd_required_skill_groups or [],
                "preferred_skill_groups": [],
                "experience_years": None,
                "experience_evidence": None,
                "education": None,
                "education_evidence": None,
                "soft_skills": [],
            },
        }

    @classmethod
    def _coverage_response(cls) -> MagicMock:
        return cls._gemini_response({"skills": [], "soft_skills": []})

    def test_extracts_structured_keywords_for_weighted_matching(self) -> None:
        output = {
            "cv": {
                "skills": [
                    {"name": "splunk", "evidence": "Splunk, Wireshark, and Python"},
                    {"name": "Wireshark", "evidence": "Splunk, Wireshark, and Python"},
                    {"name": "Python", "evidence": "Splunk, Wireshark, and Python"},
                    {"name": "Invented Skill", "evidence": "not present in the CV"},
                ],
                "experience_years": None,
                "experience_evidence": None,
                "education": "Bachelor",
                "education_evidence": "Bachelor student",
                "education_entries": [],
                "experience_entries": [],
                "soft_skills": [
                    {"name": "Communication", "evidence": "Communication"}
                ],
            },
            "jd": {
                "required_skills": [
                    {"name": "Splunk", "evidence": "Requires Splunk, Wireshark, Python"},
                    {"name": "wireshark", "evidence": "Requires Splunk, Wireshark, Python"},
                    {"name": "Python", "evidence": "Requires Splunk, Wireshark, Python"},
                ],
                "preferred_skills": [
                    {"name": "Nessus", "evidence": "Nessus preferred"}
                ],
                "required_skill_groups": [],
                "preferred_skill_groups": [],
                "experience_years": None,
                "experience_evidence": None,
                "education": "Bachelor",
                "education_evidence": "Bachelor required",
                "soft_skills": [
                    {"name": "communication", "evidence": "communication required"}
                ],
            },
        }
        response = MagicMock(status_code=200)
        response.json.return_value = {
            "candidates": [
                {
                    "content": {
                        "parts": [
                            {"text": f"```json\n{json.dumps(output)}\n```"}
                        ]
                    },
                    "finishReason": "STOP",
                }
            ]
        }

        with patch(
            "app.services.gemini_analyzer.requests.post", return_value=response
        ) as post:
            cv, jd = extract_match_inputs(
                cv_text="""Jane Doe
jane@example.com | +84 901 234 567
Bachelor student using Splunk, Wireshark, and Python. Communication.""",
                job_description="Requires Splunk, Wireshark, Python. Nessus preferred. Bachelor required; communication required.",
            )

        self.assertEqual(
            post.call_args.args[0],
            "https://generativelanguage.googleapis.com/v1beta/models/gemini-3.6-flash:generateContent",
        )
        request_body = post.call_args.kwargs["json"]
        self.assertEqual(post.call_args.kwargs["headers"]["x-goog-api-key"], "test-key")
        self.assertEqual(post.call_args.kwargs["timeout"], 1)
        self.assertEqual(
            request_body["generationConfig"]["responseMimeType"],
            "application/json",
        )
        self.assertEqual(
            request_body["generationConfig"]["thinkingConfig"],
            {"thinkingLevel": "low"},
        )
        self.assertEqual(
            request_body["generationConfig"]["maxOutputTokens"], 24_000
        )
        self.assertNotIn("temperature", request_body["generationConfig"])
        self.assertIn(
            "cv", request_body["generationConfig"]["responseJsonSchema"]["properties"]
        )
        jd_schema = request_body["generationConfig"]["responseJsonSchema"][
            "properties"
        ]["jd"]
        self.assertIn("required_skill_groups", jd_schema["properties"])
        group_skills_schema = jd_schema["properties"]["required_skill_groups"][
            "items"
        ]["properties"]["skills"]
        self.assertNotIn("minItems", group_skills_schema)
        self.assertNotIn("maxItems", group_skills_schema)
        submitted = json.loads(request_body["contents"][0]["parts"][0]["text"])[
            "cv_text"
        ]
        self.assertNotIn("Jane Doe", submitted)
        self.assertNotIn("jane@example.com", submitted)
        self.assertNotIn("+84 901 234 567", submitted)
        self.assertEqual(cv["skills"], ["Python", "splunk", "Wireshark"])
        self.assertEqual(jd["required_skills"], ["Python", "splunk", "Wireshark"])
        self.assertEqual(match_documents(cv, jd)["match_label"], "Strong Match")

    def test_extracts_and_scores_one_of_requirement_group(self) -> None:
        jd_quote = "Know at least one of C++, Python, C#, or Java."
        output = {
            "cv": {
                "skills": [{"name": "C++", "evidence": "C++ development"}],
                "experience_years": None,
                "experience_evidence": None,
                "education": None,
                "education_evidence": None,
                "education_entries": [],
                "experience_entries": [],
                "soft_skills": [],
            },
            "jd": {
                "required_skills": [],
                "preferred_skills": [],
                "required_skill_groups": [
                    {
                        "skills": [
                            {"name": skill, "evidence": jd_quote}
                            for skill in ["C++", "Python", "C#", "Java"]
                        ],
                        "minimum_required": 1,
                        "evidence": jd_quote,
                    }
                ],
                "preferred_skill_groups": [],
                "experience_years": None,
                "experience_evidence": None,
                "education": None,
                "education_evidence": None,
                "soft_skills": [],
            },
        }
        response = MagicMock(status_code=200)
        response.json.return_value = {
            "candidates": [
                {
                    "content": {"parts": [{"text": json.dumps(output)}]},
                    "finishReason": "STOP",
                }
            ]
        }

        with patch(
            "app.services.gemini_analyzer.requests.post", return_value=response
        ):
            cv, jd = extract_match_inputs(
                cv_text="Projects include C++ development and algorithms.",
                job_description=jd_quote,
            )

        self.assertEqual(jd["required_skills"], [])
        self.assertEqual(jd["required_skill_groups"][0]["minimum_required"], 1)
        result = match_documents(cv, jd)
        self.assertEqual(result["breakdown"]["skills"]["score"], 100.0)
        self.assertEqual(result["breakdown"]["skills"]["missing"], [])

    def test_analyzer_evidence_at_300_characters_needs_no_validation_retry(self) -> None:
        evidence = "A" * 300
        output = self._match_output(
            cv_skills=[{"name": "Boundary Skill", "evidence": evidence}]
        )

        with patch(
            "app.services.gemini_analyzer.requests.post",
            return_value=self._gemini_response(output),
        ) as post:
            cv, _ = extract_match_inputs(
                cv_text=evidence,
                job_description="No candidate requirements are stated.",
            )

        self.assertEqual(post.call_count, 1)
        self.assertEqual(cv["skills"], ["Boundary Skill"])
        evidence_schema = post.call_args.kwargs["json"]["generationConfig"][
            "responseJsonSchema"
        ]["properties"]["cv"]["properties"]["skills"]["items"]["properties"][
            "evidence"
        ]
        self.assertIn("at most 300 characters", evidence_schema["description"])
        self.assertNotIn("maxLength", evidence_schema)

    def test_each_required_analyzer_list_triggers_one_corrective_retry(self) -> None:
        required_paths = (
            ("cv", "education_entries"),
            ("cv", "experience_entries"),
            ("jd", "required_skill_groups"),
            ("jd", "preferred_skill_groups"),
        )
        for section, field in required_paths:
            with self.subTest(path=f"{section}.{field}"):
                invalid_output = self._match_output()
                del invalid_output[section][field]
                corrected_output = self._match_output()

                with patch(
                    "app.services.gemini_analyzer.requests.post",
                    side_effect=[
                        self._gemini_response(invalid_output),
                        self._gemini_response(corrected_output),
                    ],
                ) as post:
                    extract_match_inputs(
                        cv_text="Python backend experience.",
                        job_description="Python is required.",
                    )

                self.assertEqual(post.call_count, 2)
                correction = post.call_args_list[1].kwargs["json"]["contents"][0][
                    "parts"
                ][-1]["text"]
                self.assertIn(f"{section}.{field}:missing", correction)

    def test_validation_locations_redact_model_controlled_keys_and_values(self) -> None:
        private_key = "Jane Doe jane@example.com PRIVATE SOURCE TEXT"
        private_value = "RAW PRIVATE PAYLOAD"
        invalid_output = self._match_output(
            cv_skills=[
                {
                    "name": "Python",
                    "evidence": "Python",
                    private_key: private_value,
                }
            ]
        )

        with patch(
            "app.services.gemini_analyzer.requests.post",
            side_effect=[
                self._gemini_response(invalid_output),
                self._gemini_response(invalid_output),
            ],
        ) as post:
            with self.assertRaises(GeminiAnalyzerError) as raised:
                extract_match_inputs(
                    cv_text="Python appears in the source CV.",
                    job_description="Python is required.",
                )

        self.assertEqual(post.call_count, 2)
        correction = post.call_args_list[1].kwargs["json"]["contents"][0]["parts"][-1][
            "text"
        ]
        message = str(raised.exception)
        for rendered in (correction, message):
            self.assertIn(
                "cv.skills.0.unexpected_field:extra_forbidden",
                rendered,
            )
            self.assertNotIn(private_key, rendered)
            self.assertNotIn(private_value, rendered)
            self.assertNotIn("Jane Doe", rendered)
            self.assertNotIn("jane@example.com", rendered)
            self.assertNotIn("PRIVATE SOURCE TEXT", rendered)

    def test_search_profile_required_key_triggers_one_corrective_retry(self) -> None:
        invalid_output = {
            "job_title": "Backend Engineer",
            "keywords": ["Python"],
            "location_hint": None,
            "level": None,
        }
        corrected_output = {
            **invalid_output,
            "level_evidence": None,
        }

        with patch(
            "app.services.gemini_analyzer.requests.post",
            side_effect=[
                self._gemini_response(invalid_output),
                self._gemini_response(corrected_output),
            ],
        ) as post:
            profile = extract_cv_search_profile(cv_text="Python backend developer")

        self.assertEqual(post.call_count, 2)
        self.assertEqual(profile["job_title"], "Backend Engineer")
        correction = post.call_args_list[1].kwargs["json"]["contents"][0]["parts"][-1][
            "text"
        ]
        self.assertIn("level_evidence:missing", correction)

    def test_analyzer_evidence_at_301_characters_retries_once(self) -> None:
        invalid_evidence = "X" * 301
        corrected_evidence = "Y" * 300
        invalid_output = self._match_output(
            cv_skills=[{"name": "Corrected Skill", "evidence": invalid_evidence}]
        )
        corrected_output = self._match_output(
            cv_skills=[{"name": "Corrected Skill", "evidence": corrected_evidence}]
        )

        with patch(
            "app.services.gemini_analyzer.requests.post",
            side_effect=[
                self._gemini_response(invalid_output),
                self._gemini_response(corrected_output),
            ],
        ) as post:
            cv, _ = extract_match_inputs(
                cv_text=f"{invalid_evidence}\n{corrected_evidence}",
                job_description="No candidate requirements are stated.",
            )

        self.assertEqual(post.call_count, 2)
        self.assertEqual(cv["skills"], ["Corrected Skill"])
        correction = post.call_args_list[1].kwargs["json"]["contents"][0]["parts"][-1][
            "text"
        ]
        self.assertIn("cv.skills.0.evidence:string_too_long", correction)
        self.assertNotIn(invalid_evidence, correction)

    def test_analyzer_retry_recovers_nested_grounded_evidence(self) -> None:
        cv_evidence = (
            "Python delivery evidence. "
            + "Built reliable backend services with Python " * 10
        ).strip()
        jd_group_evidence = (
            "Know one of Python or Java. "
            + "Either language is acceptable for backend delivery " * 10
        ).strip()
        self.assertGreater(len(cv_evidence), 300)
        self.assertGreater(len(jd_group_evidence), 300)
        output = self._match_output(
            cv_skills=[{"name": "Python", "evidence": cv_evidence}],
            jd_required_skill_groups=[
                {
                    "skills": [
                        {"name": "Python", "evidence": "Python"},
                        {"name": "Java", "evidence": "Java"},
                    ],
                    "minimum_required": 1,
                    "evidence": jd_group_evidence,
                }
            ],
        )

        with patch(
            "app.services.gemini_analyzer.requests.post",
            side_effect=[self._gemini_response(output), self._gemini_response(output)],
        ) as post:
            cv, jd = extract_match_inputs(
                cv_text=cv_evidence,
                job_description=jd_group_evidence,
            )

        self.assertEqual(post.call_count, 2)
        self.assertEqual(cv["skills"], ["Python"])
        self.assertEqual(len(jd["required_skill_groups"]), 1)
        recovered = jd["required_skill_groups"][0]["evidence"]
        self.assertLessEqual(len(recovered), 300)
        self.assertIn(recovered, jd_group_evidence)

    def test_analyzer_non_repairable_failure_is_sanitized_after_one_retry(self) -> None:
        private_name = "PRIVATE-CANDIDATE-DATA-" * 7
        private_evidence = "PRIVATE-EVIDENCE-" * 30
        invalid_output = self._match_output(
            cv_skills=[{"name": private_name, "evidence": private_evidence}]
        )

        with patch(
            "app.services.gemini_analyzer.requests.post",
            side_effect=[
                self._gemini_response(invalid_output),
                self._gemini_response(invalid_output),
            ],
        ) as post:
            with self.assertRaises(GeminiAnalyzerError) as raised:
                extract_match_inputs(
                    cv_text="Source without private model output.",
                    job_description="No candidate requirements are stated.",
                )

        self.assertEqual(post.call_count, 2)
        message = str(raised.exception)
        self.assertIn("cv.skills.0.name:string_too_long", message)
        self.assertIn("cv.skills.0.evidence:string_too_long", message)
        self.assertNotIn(private_name, message)
        self.assertNotIn(private_evidence, message)
        correction = post.call_args_list[1].kwargs["json"]["contents"][0]["parts"][-1][
            "text"
        ]
        self.assertNotIn(private_name, correction)
        self.assertNotIn(private_evidence, correction)

    def test_analyzer_strict_type_failure_remains_non_repairable(self) -> None:
        invalid_output = self._match_output()
        invalid_output["cv"]["experience_years"] = "3"

        with patch(
            "app.services.gemini_analyzer.requests.post",
            side_effect=[
                self._gemini_response(invalid_output),
                self._gemini_response(invalid_output),
            ],
        ) as post:
            with self.assertRaises(GeminiAnalyzerError) as raised:
                extract_match_inputs(
                    cv_text="Three years building backend services.",
                    job_description="No candidate requirements are stated.",
                )

        self.assertEqual(post.call_count, 2)
        self.assertIn("cv.experience_years:float_type", str(raised.exception))

    def test_retries_rate_limit_once(self) -> None:
        output = {
            "cv": {
                "skills": [],
                "experience_years": None,
                "experience_evidence": None,
                "education": None,
                "education_evidence": None,
                "education_entries": [],
                "experience_entries": [],
                "soft_skills": [],
            },
            "jd": {
                "required_skills": [],
                "preferred_skills": [],
                "required_skill_groups": [],
                "preferred_skill_groups": [],
                "experience_years": None,
                "experience_evidence": None,
                "education": None,
                "education_evidence": None,
                "soft_skills": [],
            },
        }
        response = MagicMock(status_code=200)
        response.json.return_value = {
            "candidates": [
                {
                    "content": {"parts": [{"text": json.dumps(output)}]},
                    "finishReason": "STOP",
                }
            ]
        }
        rate_limit = MagicMock(status_code=429)

        with (
            patch(
                "app.services.gemini_analyzer.requests.post",
                side_effect=[rate_limit, response],
            ) as post,
            patch("app.services.gemini_analyzer.time.sleep") as retry_sleep,
        ):
            extract_match_inputs(
                cv_text="Readable CV text",
                job_description="Readable job description text",
            )

        self.assertEqual(post.call_count, 2)
        retry_sleep.assert_called_once_with(0.5)

    def test_requires_server_side_api_key(self) -> None:
        settings.gemini_api_key = None
        with self.assertRaisesRegex(GeminiAnalyzerError, "GEMINI_API_KEY"):
            extract_match_inputs(cv_text="Readable CV text", job_description="Readable job description text")

    def test_selects_gemini_model_and_new_cache_version(self) -> None:
        algorithm_version, model_name = _selected_analyzer_config()
        self.assertTrue(algorithm_version.startswith("fitcv-gemini-"))
        self.assertLessEqual(len(algorithm_version), 50)
        self.assertTrue(algorithm_version.endswith("-v8-s7"))
        self.assertEqual(model_name, "gemini-3.6-flash")

    def test_extracts_structured_cv_from_original_file(self) -> None:
        output = {
            "skills": [
                {"name": "Git", "evidence": "Git and GitHub"},
                {"name": "ASP.NET", "evidence": "ASP.NET Core"},
            ],
            "experience_years": 0.16,
            "experience_evidence": "Professional Experience",
            "education": "Bachelor",
            "education_evidence": "Major: Information Technology - Saigon Technology University",
            "education_entries": [
                {
                    "name": "Information Technology — Saigon Technology University",
                    "evidence": "Major: Information Technology - Saigon Technology University",
                }
            ],
            "experience_entries": [
                {
                    "name": "University Graduation Internship — PDF to ICS Web Application",
                    "evidence": "University Graduation Internship",
                }
            ],
            "soft_skills": [
                {"name": "Communication", "evidence": "Excellent communication"}
            ],
        }
        response = MagicMock(status_code=200)
        response.json.return_value = {
            "candidates": [
                {
                    "content": {"parts": [{"text": json.dumps(output)}]},
                    "finishReason": "STOP",
                }
            ]
        }
        coverage_response = MagicMock(status_code=200)
        coverage_response.json.return_value = {
            "candidates": [
                {
                    "content": {
                        "parts": [
                            {
                                "text": json.dumps(
                                    {
                                        "skills": [
                                            {"name": "Git", "evidence": "Git and GitHub"}
                                        ],
                                        "soft_skills": [],
                                    }
                                )
                            }
                        ]
                    },
                    "finishReason": "STOP",
                }
            ]
        }

        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            path.write_bytes(b"%PDF-fake-cv")
            with patch(
                "app.services.gemini_analyzer.requests.post",
                side_effect=[response, coverage_response],
            ) as post:
                payload = extract_cv_inputs_from_file(
                    file_path=path,
                    file_type="PDF",
                )

        request_body = post.call_args_list[0].kwargs["json"]
        self.assertEqual(
            request_body["generationConfig"]["thinkingConfig"],
            {"thinkingLevel": "low"},
        )
        self.assertEqual(
            request_body["generationConfig"]["maxOutputTokens"], 24_000
        )
        parts = request_body["contents"][0]["parts"]
        self.assertEqual(parts[0]["inlineData"]["mimeType"], "application/pdf")
        self.assertEqual(parts[0]["inlineData"]["data"], base64.b64encode(b"%PDF-fake-cv").decode("ascii"))
        self.assertEqual(payload["skills"], ["ASP.NET", "Git"])
        self.assertEqual(payload["soft_skills"], ["Communication"])
        self.assertEqual(payload["experience_years"], None)
        self.assertEqual(payload["education"], None)
        self.assertEqual(payload["education_evidence"], "Major: Information Technology - Saigon Technology University")
        self.assertEqual(
            payload["education_entries"],
            ["Information Technology — Saigon Technology University"],
        )
        self.assertEqual(
            payload["experience_entries"],
            ["University Graduation Internship — PDF to ICS Web Application"],
        )
        self.assertEqual(payload["_extraction_provider"], "gemini")

    def test_cv_evidence_at_300_characters_passes_without_validation_retry(self) -> None:
        evidence = "E" * 300
        output = self._cv_file_output(
            education_entries=[{"name": "Exact boundary", "evidence": evidence}]
        )

        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            path.write_bytes(b"%PDF-boundary")
            with patch(
                "app.services.gemini_analyzer.requests.post",
                side_effect=[
                    self._gemini_response(output),
                    self._coverage_response(),
                ],
            ) as post:
                payload = extract_cv_inputs_from_file(
                    file_path=path,
                    file_type="PDF",
                    source_text=evidence,
                )

        self.assertEqual(post.call_count, 2)
        self.assertEqual(payload["education_entries"], ["Exact boundary"])
        evidence_schema = post.call_args_list[0].kwargs["json"]["generationConfig"][
            "responseJsonSchema"
        ]["properties"]["education_entries"]["items"]["properties"]["evidence"]
        self.assertIn("at most 300 characters", evidence_schema["description"])
        self.assertNotIn("maxLength", evidence_schema)

    def test_cv_evidence_at_301_characters_retries_once_and_accepts_correction(self) -> None:
        invalid_evidence = "X" * 301
        corrected_evidence = "Y" * 300
        invalid_output = self._cv_file_output(
            education_entries=[
                {"name": "Corrected education", "evidence": invalid_evidence}
            ]
        )
        corrected_output = self._cv_file_output(
            education_entries=[
                {"name": "Corrected education", "evidence": corrected_evidence}
            ]
        )

        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            path.write_bytes(b"%PDF-retry")
            with patch(
                "app.services.gemini_analyzer.requests.post",
                side_effect=[
                    self._gemini_response(invalid_output),
                    self._gemini_response(corrected_output),
                    self._coverage_response(),
                ],
            ) as post:
                payload = extract_cv_inputs_from_file(
                    file_path=path,
                    file_type="PDF",
                    source_text=f"{invalid_evidence} {corrected_evidence}",
                )

        self.assertEqual(post.call_count, 3)
        self.assertEqual(payload["education_entries"], ["Corrected education"])
        correction = post.call_args_list[1].kwargs["json"]["contents"][0]["parts"][-1][
            "text"
        ]
        self.assertIn("education_entries.0.evidence:string_too_long", correction)
        self.assertNotIn(invalid_evidence, correction)

    def test_retry_falls_back_to_grounded_word_bounded_evidence_excerpt(self) -> None:
        long_evidence = (
            "Bachelor degree in Computer Science. "
            + "Completed advanced software engineering coursework " * 8
        ).strip()
        self.assertGreater(len(long_evidence), 300)
        output = self._cv_file_output(
            education="Bachelor",
            education_evidence=long_evidence,
            education_entries=[
                {"name": "Bachelor degree", "evidence": long_evidence}
            ],
        )

        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            path.write_bytes(b"%PDF-grounded-fallback")
            with patch(
                "app.services.gemini_analyzer.requests.post",
                side_effect=[
                    self._gemini_response(output),
                    self._gemini_response(output),
                    self._coverage_response(),
                ],
            ) as post:
                payload = extract_cv_inputs_from_file(
                    file_path=path,
                    file_type="PDF",
                    source_text=long_evidence,
                )

        self.assertEqual(post.call_count, 3)
        self.assertEqual(payload["education"], "Bachelor")
        self.assertEqual(payload["education_entries"], ["Bachelor degree"])
        self.assertLessEqual(len(payload["education_evidence"]), 300)
        self.assertIn(payload["education_evidence"], long_evidence)
        self.assertFalse(payload["education_evidence"].endswith(" "))

    def test_retry_drops_ungrounded_overlong_optional_entry(self) -> None:
        ungrounded_evidence = "Invented private evidence " * 20
        output = self._cv_file_output(
            education="Bachelor",
            education_evidence=ungrounded_evidence,
            education_entries=[
                {"name": "Unverified education", "evidence": ungrounded_evidence}
            ]
        )

        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            path.write_bytes(b"%PDF-ungrounded-fallback")
            with patch(
                "app.services.gemini_analyzer.requests.post",
                side_effect=[
                    self._gemini_response(output),
                    self._gemini_response(output),
                    self._coverage_response(),
                ],
            ) as post:
                payload = extract_cv_inputs_from_file(
                    file_path=path,
                    file_type="PDF",
                    source_text="Bachelor of Science in Computer Science.",
                )

        self.assertEqual(post.call_count, 3)
        self.assertEqual(payload["education_entries"], [])
        self.assertIsNone(payload["education_evidence"])
        self.assertIsNone(payload["education"])

    def test_final_cv_validation_error_contains_only_sanitized_issues(self) -> None:
        private_name = "PRIVATE-CANDIDATE-DATA-" * 7
        invalid_output = self._cv_file_output(
            education_entries=[
                {"name": private_name, "evidence": "Z" * 301}
            ]
        )

        with TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            path.write_bytes(b"%PDF-invalid")
            with patch(
                "app.services.gemini_analyzer.requests.post",
                side_effect=[
                    self._gemini_response(invalid_output),
                    self._gemini_response(invalid_output),
                ],
            ) as post:
                with self.assertRaises(GeminiAnalyzerError) as raised:
                    extract_cv_inputs_from_file(
                        file_path=path,
                        file_type="PDF",
                        source_text="Source without the private model output.",
                    )

        self.assertEqual(post.call_count, 2)
        message = str(raised.exception)
        self.assertIn("education_entries.0.name:string_too_long", message)
        self.assertIn("education_entries.0.evidence:string_too_long", message)
        self.assertNotIn(private_name, message)


class AnalyzerRepositoryTests(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = create_engine("sqlite+pysqlite:///:memory:")
        Base.metadata.create_all(
            self.engine,
            tables=[
                Account.__table__,
                Cv.__table__,
                CvParseResult.__table__,
                Job.__table__,
                JobDescription.__table__,
                JdParseResult.__table__,
                MatchResult.__table__,
                AiTask.__table__,
            ],
        )
        self.db = Session(self.engine)
        self.account = Account(
            email="student@example.com",
            password_hash="test",
            full_name="Student",
            auth_provider=AuthProvider.password,
        )
        self.db.add(self.account)
        self.db.commit()
        self.db.refresh(self.account)

    def tearDown(self) -> None:
        self.db.close()
        self.engine.dispose()

    def test_versions_dedupe_and_persist_match_evidence(self) -> None:
        first, _ = analyzer.create_cv(
            self.db,
            account_id=self.account.account_id,
            file_name="cv-v1.pdf",
            file_path="cv/1/v1.pdf",
            file_type="PDF",
            file_size_kb=10,
            file_sha256="1" * 64,
            parser_version=PARSER_VERSION,
        )
        second, parsed_cv = analyzer.create_cv(
            self.db,
            account_id=self.account.account_id,
            file_name="cv-v2.docx",
            file_path="cv/1/v2.docx",
            file_type="DOCX",
            file_size_kb=20,
            file_sha256="2" * 64,
            parser_version=PARSER_VERSION,
        )
        self.assertEqual((first.version_number, second.version_number), (1, 2))
        self.assertFalse(first.is_latest)
        self.assertTrue(second.is_latest)

        cv_payload = {"skills": ["Python"], "experience_years": 2, "education": None, "soft_skills": []}
        jd_payload = {"required_skills": ["Python"], "preferred_skills": [], "experience_years": 2, "education": None, "soft_skills": []}
        analyzer.set_parse_success(self.db, parsed_cv, text="Python developer with 2 years experience", payload=cv_payload)
        description, parsed_jd = analyzer.get_or_create_job_description(
            self.db,
            account_id=self.account.account_id,
            title="Python Developer",
            raw_text="Python developer with 2 years experience required for this backend role.",
            content_sha256="3" * 64,
            parsed_payload=jd_payload,
            parser_version=PARSER_VERSION,
        )
        duplicate, duplicate_parse = analyzer.get_or_create_job_description(
            self.db,
            account_id=self.account.account_id,
            title="Python Developer",
            raw_text="Python developer with 2 years experience required for this backend role.",
            content_sha256="3" * 64,
            parsed_payload=jd_payload,
            parser_version=PARSER_VERSION,
        )
        self.assertEqual((description.job_description_id, parsed_jd.jd_parse_id), (duplicate.job_description_id, duplicate_parse.jd_parse_id))

        match = analyzer.create_pending_match(
            self.db,
            cv=second,
            parsed_cv=parsed_cv,
            description=description,
            parsed_jd=parsed_jd,
            algorithm_version=ALGORITHM_VERSION,
        )
        analyzer.set_match_success(self.db, match, match_documents(cv_payload, jd_payload))
        self.assertEqual(match.status, "Success")
        self.assertEqual(float(match.overall_score or 0), 100.0)
        self.assertIsNotNone(match.evidence_json)

    def test_cv_parse_stays_processing_until_terminal_queue_attempt(self) -> None:
        cv, _ = analyzer.create_cv(
            self.db,
            account_id=self.account.account_id,
            file_name="cv.pdf",
            file_path="cv/1/cv.pdf",
            file_type="PDF",
            file_size_kb=10,
            file_sha256="4" * 64,
            parser_version=PARSER_VERSION,
        )
        session_factory = sessionmaker(bind=self.engine)
        parse_error = GeminiAnalyzerError(
            "Gemini returned invalid CV extraction data "
            "(education_entries.0.evidence: string_too_long)."
        )

        with (
            patch("app.services.analyzer_service.SessionLocal", session_factory),
            patch(
                "app.services.analyzer_service.extract_document_text",
                return_value="CV source text",
            ),
            patch(
                "app.services.analyzer_service.selected_analyzer_config",
                return_value=("fitcv-gemini-test", "gemini-test"),
            ),
            patch(
                "app.services.analyzer_service.extract_cv_inputs_from_file",
                side_effect=parse_error,
            ),
        ):
            self.assertFalse(
                run_cv_parse(cv.cv_id, terminal_failure=False)
            )
            self.db.expire_all()
            parsed = analyzer.get_latest_parse(self.db, cv.cv_id)
            self.assertIsNotNone(parsed)
            self.assertEqual(parsed.parse_status, "Processing")
            self.assertIsNone(parsed.error_message)

            self.assertFalse(run_cv_parse(cv.cv_id, terminal_failure=True))
            self.db.expire_all()
            parsed = analyzer.get_latest_parse(self.db, cv.cv_id)
            self.assertIsNotNone(parsed)
            self.assertEqual(parsed.parse_status, "Failed")
            self.assertIn("education_entries.0.evidence", parsed.error_message)


class AnalyzerApiTests(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = create_engine(
            "sqlite+pysqlite:///:memory:",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        self.session_factory = sessionmaker(bind=self.engine)
        Base.metadata.create_all(
            self.engine,
            tables=[
                Account.__table__,
                Cv.__table__,
                CvParseResult.__table__,
                Job.__table__,
                JobDescription.__table__,
                JdParseResult.__table__,
                MatchResult.__table__,
                AiTask.__table__,
            ],
        )
        db = self.session_factory()
        self.account = Account(
            email="api-student@example.com",
            password_hash="test",
            full_name="API Student",
            auth_provider=AuthProvider.password,
        )
        db.add(self.account)
        db.commit()
        db.refresh(self.account)
        db.expunge(self.account)
        db.close()

        def override_db():
            session = self.session_factory()
            try:
                yield session
            finally:
                session.close()

        self.uploads = TemporaryDirectory()
        self.original_upload_dir = settings.upload_dir
        self.original_analyzer_provider = settings.analyzer_provider
        settings.upload_dir = Path(self.uploads.name)
        settings.analyzer_provider = "deterministic"
        app.dependency_overrides[get_db] = override_db
        app.dependency_overrides[get_current_account] = lambda: self.account
        self.session_patch = patch("app.services.analyzer_service.SessionLocal", self.session_factory)
        self.session_patch.start()
        self.client = TestClient(app)

    def tearDown(self) -> None:
        self.client.close()
        self.session_patch.stop()
        app.dependency_overrides.clear()
        settings.upload_dir = self.original_upload_dir
        settings.analyzer_provider = self.original_analyzer_provider
        self.uploads.cleanup()
        self.engine.dispose()

    def test_failed_cv_parse_can_enqueue_a_new_idempotent_retry(self) -> None:
        document = Document()
        document.add_heading("Technical Skills")
        document.add_paragraph("Python, FastAPI, MySQL and communication")
        document.add_heading("Experience")
        document.add_paragraph("3 years building REST APIs.")
        buffer = BytesIO()
        document.save(buffer)

        upload = self.client.post(
            "/api/cvs",
            files={
                "file": (
                    "resume.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document",
                )
            },
        )
        self.assertEqual(upload.status_code, 201, upload.text)
        cv_id = upload.json()["cv_id"]
        original_task_id = upload.json()["ai_task_id"]

        db = self.session_factory()
        try:
            parsed = analyzer.get_latest_parse(db, cv_id)
            self.assertIsNotNone(parsed)
            analyzer.set_parse_failed(db, parsed, "Previous terminal parse failure.")
        finally:
            db.close()

        retried = self.client.post(f"/api/cvs/{cv_id}/retry-parse")
        self.assertEqual(retried.status_code, 202, retried.text)
        self.assertNotEqual(retried.json()["ai_task_id"], original_task_id)

        completed = self.client.get(f"/api/cvs/{cv_id}")
        self.assertEqual(completed.status_code, 200, completed.text)
        self.assertEqual(completed.json()["parse_status"], "Success")
        self.assertIsNone(completed.json()["error_message"])

        already_complete = self.client.post(f"/api/cvs/{cv_id}/retry-parse")
        self.assertEqual(already_complete.status_code, 409, already_complete.text)

    def test_upload_analyze_history_and_delete(self) -> None:
        document = Document()
        document.add_heading("Technical Skills")
        document.add_paragraph("Python, FastAPI, MySQL, Docker and communication")
        document.add_heading("Experience")
        document.add_paragraph("3 years building REST APIs.")
        buffer = BytesIO()
        document.save(buffer)

        upload = self.client.post(
            "/api/cvs",
            files={"file": ("resume.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        self.assertEqual(upload.status_code, 201, upload.text)
        cv_id = upload.json()["cv_id"]
        parsed = self.client.get(f"/api/cvs/{cv_id}")
        self.assertEqual(parsed.json()["parse_status"], "Success")

        analysis = self.client.post(
            "/api/analyzer/matches",
            json={
                "cv_id": cv_id,
                "job_description": "Backend role requires 2 years of Python, FastAPI, MySQL and REST API experience with strong communication skills.",
                "title": "Backend Developer",
            },
        )
        self.assertEqual(analysis.status_code, 202, analysis.text)
        match_id = analysis.json()["match_result_id"]
        completed = self.client.get(f"/api/analyzer/matches/{match_id}")
        self.assertEqual(completed.json()["status"], "Success")
        self.assertIn(completed.json()["match_label"], {"Strong Match", "Moderate Match", "Weak Match"})
        improved = Document()
        improved.add_heading("Technical Skills")
        improved.add_paragraph("Python, FastAPI, MySQL, Docker, Kubernetes and communication")
        improved.add_heading("Experience")
        improved.add_paragraph("5 years building REST APIs and cloud services.")
        improved_buffer = BytesIO()
        improved.save(improved_buffer)
        second_upload = self.client.post(
            "/api/cvs",
            files={
                "file": (
                    "resume-v2.docx",
                    improved_buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(second_upload.status_code, 201, second_upload.text)
        second_cv_id = second_upload.json()["cv_id"]
        comparison = self.client.get(
            f"/api/cvs/compare?base_cv_id={cv_id}&target_cv_id={second_cv_id}"
        )
        self.assertEqual(comparison.status_code, 200, comparison.text)
        self.assertEqual(comparison.json()["base"]["version_number"], 1)
        self.assertEqual(comparison.json()["target"]["version_number"], 2)
        skills_change = next(
            item for item in comparison.json()["changes"] if item["category"] == "Skills"
        )
        self.assertIn("Kubernetes", skills_change["added"])
        self.assertEqual(len(self.client.get("/api/cvs").json()), 2)
        self.assertEqual(self.client.delete(f"/api/cvs/{cv_id}").status_code, 204)
        self.assertEqual(len(self.client.get("/api/cvs").json()), 1)

    def test_failed_match_enqueues_a_new_idempotent_retry(self) -> None:
        document = Document()
        document.add_heading("Technical Skills")
        document.add_paragraph("Python, FastAPI, MySQL and communication")
        document.add_heading("Experience")
        document.add_paragraph("3 years building REST APIs.")
        buffer = BytesIO()
        document.save(buffer)
        request = {
            "job_description": (
                "Backend role requires 2 years of Python, FastAPI, MySQL and "
                "REST API experience with strong communication skills."
            ),
            "title": "Backend Developer",
        }

        upload = self.client.post(
            "/api/cvs",
            files={
                "file": (
                    "resume.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document",
                )
            },
        )
        self.assertEqual(upload.status_code, 201, upload.text)
        request["cv_id"] = upload.json()["cv_id"]

        first = self.client.post("/api/analyzer/matches", json=request)
        self.assertEqual(first.status_code, 202, first.text)
        match_id = first.json()["match_result_id"]

        db = self.session_factory()
        try:
            first_task = ai_tasks.get_latest_for_resource(
                db, task_type="MatchAnalysis", resource_id=match_id
            )
            self.assertIsNotNone(first_task)
            first_task_id = first_task.ai_task_id
            match = db.get(MatchResult, match_id)
            self.assertIsNotNone(match)
            analyzer.set_match_failed(db, match, "Previous terminal match failure.")
        finally:
            db.close()

        retried = self.client.post("/api/analyzer/matches", json=request)
        self.assertEqual(retried.status_code, 202, retried.text)

        db = self.session_factory()
        try:
            retry_task = ai_tasks.get_latest_for_resource(
                db, task_type="MatchAnalysis", resource_id=match_id
            )
            self.assertIsNotNone(retry_task)
            self.assertNotEqual(retry_task.ai_task_id, first_task_id)
            self.assertEqual(
                retry_task.idempotency_key,
                f"match-analysis:{match_id}:retry:{first_task_id}",
            )
        finally:
            db.close()

        completed = self.client.get(f"/api/analyzer/matches/{match_id}")
        self.assertEqual(completed.status_code, 200, completed.text)
        self.assertEqual(completed.json()["status"], "Success")


if __name__ == "__main__":
    unittest.main()
