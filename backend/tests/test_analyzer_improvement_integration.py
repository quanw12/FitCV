from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import func, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool
from sqlalchemy import create_engine

from app.api.deps import get_current_account
from app.core.config import settings
from app.db.session import Base, get_db
from app.main import app
from app.models.account import Account, AccountRole, AuthProvider
from app.models.analyzer import Cv, CvParseResult, JdParseResult, Job, JobDescription, MatchResult
from app.models.improvement import AiTask, AiTaskStatus, CvImprovementSuggestion
from app.repositories import improvements as improvements_repository
from app.schemas.improvement import ImprovementReportData
from app.services.improvement_service import run_generation_task
from app.services.improvement_provider import ImprovementProviderError


@dataclass
class FlowHarness:
    client: TestClient
    session_factory: sessionmaker
    current_account: dict[str, Account]
    account: Account


class StubImprovementProvider:
    name = "stub"
    model_name = "stub-improvement-v1"

    def __init__(self, *, fail: bool = False) -> None:
        self.fail = fail
        self.calls: list[dict] = []

    def generate_improvement_report(
        self,
        *,
        parsed_cv: dict | str | None,
        job_description: str,
        match_result: dict,
        raw_cv_text: str | None = None,
    ) -> ImprovementReportData:
        self.calls.append(
            {
                "parsed_cv": parsed_cv,
                "raw_cv_text": raw_cv_text,
                "job_description": job_description,
                "match_result": match_result,
            }
        )
        if self.fail:
            raise ImprovementProviderError("Synthetic provider failure.")
        return ImprovementReportData.model_validate(
            {
                "skill_gaps": [
                    {
                        "skill": "Kubernetes",
                        "priority": "High",
                        "reason": "The analyzed CV does not show Kubernetes experience.",
                        "jd_evidence": "Kubernetes deployment experience is preferred.",
                    }
                ],
                "section_feedback": [
                    {
                        "section": "WorkExperience",
                        "issue": "The experience statement does not describe an outcome.",
                        "explanation": "A verified result would make the evidence easier to assess.",
                        "priority": "Medium",
                        "suggested_action": "Add a real, verifiable outcome or a placeholder.",
                    }
                ],
                "rewrite_suggestions": [
                    {
                        "section": "WorkExperience",
                        "original_text": "3 years building REST APIs.",
                        "issue": "The statement lacks a concrete outcome.",
                        "suggested_text": "Built REST APIs over 3 years, improving [verified outcome].",
                        "framework": "Action → Result",
                    }
                ],
                "quick_wins": [
                    {
                        "title": "Add one verified API outcome",
                        "category": "Experience",
                        "priority": "Medium",
                        "explanation": "Replace the placeholder only with a result you can prove.",
                    }
                ],
            }
        )


@pytest.fixture
def flow_harness() -> FlowHarness:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    factory = sessionmaker(bind=engine)
    Base.metadata.create_all(
        engine,
        tables=[
            Account.__table__,
            Cv.__table__,
            CvParseResult.__table__,
            Job.__table__,
            JobDescription.__table__,
            JdParseResult.__table__,
            MatchResult.__table__,
            CvImprovementSuggestion.__table__,
            AiTask.__table__,
        ],
    )
    db = factory()
    account = Account(
        email="flow-student@example.com",
        password_hash="test",
        full_name="Flow Student",
        role=AccountRole.student,
        auth_provider=AuthProvider.password,
    )
    db.add(account)
    db.commit()
    db.refresh(account)
    db.expunge(account)
    db.close()

    current_account = {"value": account}

    def override_db():
        session = factory()
        try:
            yield session
        finally:
            session.close()

    uploads = TemporaryDirectory()
    original_upload_dir = settings.upload_dir
    original_analyzer_provider = settings.analyzer_provider
    settings.upload_dir = Path(uploads.name)
    settings.analyzer_provider = "deterministic"
    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_current_account] = lambda: current_account["value"]
    analyzer_session_patch = patch("app.services.analyzer_service.SessionLocal", factory)
    improvement_session_patch = patch("app.services.improvement_service.SessionLocal", factory)
    analyzer_session_patch.start()
    improvement_session_patch.start()

    client = TestClient(app)
    try:
        yield FlowHarness(client, factory, current_account, account)
    finally:
        client.close()
        improvement_session_patch.stop()
        analyzer_session_patch.stop()
        app.dependency_overrides.clear()
        settings.upload_dir = original_upload_dir
        settings.analyzer_provider = original_analyzer_provider
        uploads.cleanup()
        engine.dispose()


def _create_completed_match(client: TestClient) -> tuple[int, int]:
    document = Document()
    document.add_heading("Technical Skills")
    document.add_paragraph("Python, FastAPI, MySQL, Docker and communication")
    document.add_heading("Experience")
    document.add_paragraph("3 years building REST APIs.")
    buffer = BytesIO()
    document.save(buffer)

    upload = client.post(
        "/api/cvs",
        files={
            "file": (
                "synthetic-resume.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 201, upload.text
    cv_id = upload.json()["cv_id"]
    assert client.get(f"/api/cvs/{cv_id}").json()["parse_status"] == "Success"

    analysis = client.post(
        "/api/analyzer/matches",
        json={
            "cv_id": cv_id,
            "job_description": (
                "Backend role requires 2 years of Python, FastAPI, MySQL and REST API "
                "experience with strong communication skills. "
                "Kubernetes deployment experience is preferred."
            ),
            "title": "Backend Developer",
        },
    )
    assert analysis.status_code == 202, analysis.text
    match_id = analysis.json()["match_result_id"]
    completed = client.get(f"/api/analyzer/matches/{match_id}")
    assert completed.status_code == 200, completed.text
    assert completed.json()["status"] == "Success"
    return cv_id, match_id


def _add_newer_unrelated_parse(harness: FlowHarness, cv_id: int) -> None:
    with harness.session_factory() as db:
        db.add(
            CvParseResult(
                cv_id=cv_id,
                parsed_text="A replacement parse mentioning Kubernetes.",
                parsed_json={
                    "skills": ["Kubernetes"],
                    "sections": {"skills": "Kubernetes"},
                },
                parse_status="Success",
                parser_version="synthetic-newer-v1",
            )
        )
        db.commit()


def _add_account(harness: FlowHarness, *, email: str, role: AccountRole) -> Account:
    with harness.session_factory() as db:
        account = Account(
            email=email,
            password_hash="test",
            full_name="Other Account",
            role=role,
            auth_provider=AuthProvider.password,
        )
        db.add(account)
        db.commit()
        db.refresh(account)
        db.expunge(account)
        return account


def test_analyzer_result_generates_and_persists_improvement_report(
    flow_harness: FlowHarness,
) -> None:
    cv_id, match_id = _create_completed_match(flow_harness.client)
    _add_newer_unrelated_parse(flow_harness, cv_id)
    provider = StubImprovementProvider()

    pending = flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    )
    assert pending.status_code == 200
    assert pending.json()["match_result_id"] == match_id
    assert pending.json()["status"] == "Pending"

    with patch(
        "app.services.improvement_service.get_improvement_provider",
        return_value=provider,
    ):
        generated = flow_harness.client.post(
            f"/api/match-results/{match_id}/improvement-report/generate"
        )

    assert generated.status_code == 202, generated.text
    assert generated.json()["match_result_id"] == match_id
    assert generated.json()["task_id"] is not None
    assert len(provider.calls) == 1
    assert provider.calls[0]["parsed_cv"]["skills"] == [
        "Docker",
        "FastAPI",
        "MySQL",
        "Python",
    ]
    assert "3 years building REST APIs." in provider.calls[0]["raw_cv_text"]
    assert isinstance(provider.calls[0]["match_result"]["strengths"], list)
    assert isinstance(provider.calls[0]["match_result"]["weaknesses"], list)

    report_response = flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    )
    assert report_response.status_code == 200, report_response.text
    payload = report_response.json()
    assert payload["match_result_id"] == match_id
    assert payload["status"] == "Success"
    assert payload["stale"] is False
    assert {item["skill"] for item in payload["report"]["skill_gaps"]} == {
        "Kubernetes",
        "REST API",
    }
    assert len(payload["report"]["section_feedback"]) >= 2
    assert len(payload["report"]["rewrite_suggestions"]) >= 1
    assert len(payload["report"]["quick_wins"]) >= 3

    duplicate = flow_harness.client.post(
        f"/api/match-results/{match_id}/improvement-report/generate"
    )
    assert duplicate.status_code == 202
    assert duplicate.json()["status"] == "Success"
    assert duplicate.json()["task_id"] == generated.json()["task_id"]

    reloaded = flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    )
    assert reloaded.json() == payload

    with flow_harness.session_factory() as db:
        task = db.scalar(
            select(AiTask)
            .where(AiTask.resource_id == match_id)
            .order_by(AiTask.ai_task_id.desc())
        )
        suggestion_count = db.scalar(
            select(func.count(CvImprovementSuggestion.suggestion_id)).where(
                CvImprovementSuggestion.match_result_id == match_id
            )
        )
        assert task is not None and task.status == AiTaskStatus.success
        report = payload["report"]
        assert suggestion_count == sum(
            len(report[collection])
            for collection in (
                "skill_gaps",
                "section_feedback",
                "rewrite_suggestions",
                "quick_wins",
            )
        )


def test_improvement_authorization_and_incomplete_match_errors(
    flow_harness: FlowHarness,
) -> None:
    _, match_id = _create_completed_match(flow_harness.client)

    account_override = app.dependency_overrides.pop(get_current_account)
    try:
        unauthenticated = flow_harness.client.get(
            f"/api/match-results/{match_id}/improvement-report"
        )
        assert unauthenticated.status_code == 401
    finally:
        app.dependency_overrides[get_current_account] = account_override

    hr = _add_account(flow_harness, email="flow-hr@example.com", role=AccountRole.hr)
    flow_harness.current_account["value"] = hr
    forbidden = flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    )
    assert forbidden.status_code == 403

    other_student = _add_account(
        flow_harness,
        email="flow-other@example.com",
        role=AccountRole.student,
    )
    flow_harness.current_account["value"] = other_student
    hidden = flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    )
    assert hidden.status_code == 404

    flow_harness.current_account["value"] = flow_harness.account
    with flow_harness.session_factory() as db:
        match = db.get(MatchResult, match_id)
        assert match is not None
        match.status = "Processing"
        match.overall_score = None
        db.commit()

    incomplete = flow_harness.client.post(
        f"/api/match-results/{match_id}/improvement-report/generate"
    )
    assert incomplete.status_code == 409

    with flow_harness.session_factory() as db:
        match = db.get(MatchResult, match_id)
        assert match is not None
        match.status = "Success"
        match.overall_score = 75
        match.job_description_id = None
        db.commit()

    missing_snapshot = flow_harness.client.post(
        f"/api/match-results/{match_id}/improvement-report/generate"
    )
    assert missing_snapshot.status_code == 409
    assert "immutable job-description snapshot" in missing_snapshot.json()["detail"]

    missing_snapshot_report = flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    )
    assert missing_snapshot_report.status_code == 409
    assert "immutable job-description snapshot" in missing_snapshot_report.json()["detail"]


def test_legacy_task_without_job_snapshot_fails_without_calling_provider(
    flow_harness: FlowHarness,
) -> None:
    _, match_id = _create_completed_match(flow_harness.client)
    with flow_harness.session_factory() as db:
        match = db.get(MatchResult, match_id)
        assert match is not None
        match.job_description_id = None
        task = improvements_repository.create_task(
            db,
            match_result_id=match_id,
            provider="stub",
            model_name="stub-improvement-v1",
        )
        db.commit()
        task_id = task.ai_task_id

    provider = StubImprovementProvider()
    with patch(
        "app.services.improvement_service.get_improvement_provider",
        return_value=provider,
    ):
        run_generation_task(task_id)

    assert provider.calls == []
    with flow_harness.session_factory() as db:
        task = db.get(AiTask, task_id)
        assert task is not None
        assert task.status == AiTaskStatus.failed
        assert task.error_message == "Improvement generation failed. Please try again."


def test_failed_improvement_task_can_be_retried(
    flow_harness: FlowHarness,
) -> None:
    _, match_id = _create_completed_match(flow_harness.client)
    failed_provider = StubImprovementProvider(fail=True)
    with patch(
        "app.services.improvement_service.get_improvement_provider",
        return_value=failed_provider,
    ):
        first = flow_harness.client.post(
            f"/api/match-results/{match_id}/improvement-report/generate"
        )
    assert first.status_code == 202

    failed = flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    )
    assert failed.json()["status"] == "Failed"
    assert failed.json()["error_message"] == "Synthetic provider failure."

    successful_provider = StubImprovementProvider()
    with patch(
        "app.services.improvement_service.get_improvement_provider",
        return_value=successful_provider,
    ):
        retried = flow_harness.client.post(
            f"/api/match-results/{match_id}/improvement-report/generate"
        )
    assert retried.status_code == 202
    assert retried.json()["task_id"] != first.json()["task_id"]
    assert flow_harness.client.get(
        f"/api/match-results/{match_id}/improvement-report"
    ).json()["status"] == "Success"
