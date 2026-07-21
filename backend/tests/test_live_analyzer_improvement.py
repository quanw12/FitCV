"""Opt-in Railway + Gemini smoke test.

Run only after rotating credentials and applying migration 004:
    FITCV_RUN_RAILWAY_E2E=1 python -m pytest tests/test_live_analyzer_improvement.py -q -s

The test uses synthetic data, exercises the public API contract, and removes its
database rows and uploaded file in a finally block. It is skipped in the normal suite.
"""

from __future__ import annotations

import os
import time
import uuid
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import delete, inspect, select

from app.core.config import settings
from app.db.session import SessionLocal, engine
from app.main import app
from app.models.account import Account
from app.models.analyzer import Cv, JobDescription, MatchResult
from app.models.improvement import AiTask, CvImprovementSuggestion


RUN_LIVE_E2E = os.getenv("FITCV_RUN_RAILWAY_E2E") == "1"


@pytest.mark.skipif(
    not RUN_LIVE_E2E,
    reason="Set FITCV_RUN_RAILWAY_E2E=1 after rotating secrets to run Railway E2E.",
)
def test_live_analyzer_to_improvement_flow_and_cleanup() -> None:
    assert settings.analyzer_provider == "gemini", (
        "Railway E2E requires ANALYZER_PROVIDER=gemini."
    )
    assert settings.gemini_api_key, "Railway E2E requires GEMINI_API_KEY."
    assert inspect(engine).has_table("ai_task"), (
        "Railway is missing ai_task; apply migration 004 before running live E2E."
    )

    unique = uuid.uuid4().hex
    email = f"fitcv-e2e-{unique}@example.com"
    password = f"FitCV-{unique}-Aa1!"
    token: str | None = None
    cv_id: int | None = None
    match_id: int | None = None
    job_description_id: int | None = None
    uploaded_path: Path | None = None

    with TestClient(app) as client:
        try:
            registered = client.post(
                "/api/auth/register",
                json={
                    "email": email,
                    "password": password,
                    "full_name": "Synthetic FitCV E2E Student",
                },
            )
            assert registered.status_code == 201, registered.text
            token = registered.json()["access_token"]
            headers = {"Authorization": f"Bearer {token}"}

            selected = client.post(
                "/api/auth/select-role",
                json={"role": "Student"},
                headers=headers,
            )
            assert selected.status_code == 200, selected.text

            document = Document()
            document.add_heading("Technical Skills")
            document.add_paragraph("Python, FastAPI, MySQL, REST APIs and communication")
            document.add_heading("Experience")
            document.add_paragraph(
                "Built backend APIs for a synthetic classroom project and documented endpoints."
            )
            document.add_heading("Education")
            document.add_paragraph("Bachelor of Computer Science student")
            buffer = BytesIO()
            document.save(buffer)

            uploaded = client.post(
                "/api/cvs",
                files={
                    "file": (
                        "synthetic-fitcv-e2e.docx",
                        buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=headers,
            )
            assert uploaded.status_code == 201, uploaded.text
            cv_id = uploaded.json()["cv_id"]
            cv_payload = _poll_json(
                client,
                f"/api/cvs/{cv_id}",
                headers,
                status_field="parse_status",
            )
            assert cv_payload["parse_status"] == "Success", cv_payload

            with SessionLocal() as db:
                cv = db.get(Cv, cv_id)
                assert cv is not None
                uploaded_path = settings.upload_dir / cv.file_path

            analyzed = client.post(
                "/api/analyzer/matches",
                json={
                    "cv_id": cv_id,
                    "title": "Synthetic Platform Engineer",
                    "job_description": (
                        "The Platform Engineer role requires Python, FastAPI, REST API and MySQL "
                        "experience. Docker and Kubernetes deployment skills are required. "
                        "CI/CD, Redis, monitoring, collaboration and clear communication are preferred."
                    ),
                },
                headers=headers,
            )
            assert analyzed.status_code == 202, analyzed.text
            match_id = analyzed.json()["match_result_id"]
            match_payload = _poll_json(
                client,
                f"/api/analyzer/matches/{match_id}",
                headers,
                status_field="status",
            )
            assert match_payload["status"] == "Success", match_payload
            job_description_id = match_payload["job_description_id"]

            generated = client.post(
                f"/api/match-results/{match_id}/improvement-report/generate",
                headers=headers,
            )
            assert generated.status_code == 202, generated.text
            assert generated.json()["match_result_id"] == match_id

            report_payload = _poll_json(
                client,
                f"/api/match-results/{match_id}/improvement-report",
                headers,
                status_field="status",
                attempts=45,
                interval_seconds=2,
            )
            assert report_payload["match_result_id"] == match_id
            assert report_payload["status"] == "Success", report_payload
            assert report_payload["overall_score"] == match_payload["overall_score"]
            report = report_payload["report"]
            assert report is not None
            for section in (
                "skill_gaps",
                "section_feedback",
                "rewrite_suggestions",
                "quick_wins",
            ):
                assert isinstance(report[section], list)
                assert report[section], f"Live report returned an empty {section}."
            assert len(report["section_feedback"]) >= 2
            assert len(report["rewrite_suggestions"]) >= 1
            assert len(report["quick_wins"]) >= 3

            reloaded = client.get(
                f"/api/match-results/{match_id}/improvement-report",
                headers=headers,
            )
            assert reloaded.status_code == 200
            assert reloaded.json()["report"] == report
        finally:
            if token and cv_id is not None:
                client.delete(
                    f"/api/cvs/{cv_id}",
                    headers={"Authorization": f"Bearer {token}"},
                )

            with SessionLocal() as db:
                if match_id is not None:
                    db.execute(
                        delete(AiTask).where(
                            AiTask.task_type == "ImprovementReport",
                            AiTask.resource_id == match_id,
                        )
                    )
                account = db.scalar(select(Account).where(Account.email == email))
                if account is not None:
                    db.delete(account)
                db.commit()

            if uploaded_path is not None:
                uploaded_path.unlink(missing_ok=True)
                assert not uploaded_path.exists(), (
                    f"Synthetic upload was not cleaned up: {uploaded_path}"
                )

    with SessionLocal() as db:
        assert db.scalar(select(Account).where(Account.email == email)) is None
        if cv_id is not None:
            assert db.get(Cv, cv_id) is None
        if match_id is not None:
            assert db.get(MatchResult, match_id) is None
            assert db.scalar(
                select(AiTask).where(
                    AiTask.task_type == "ImprovementReport",
                    AiTask.resource_id == match_id,
                )
            ) is None
            assert db.scalar(
                select(CvImprovementSuggestion).where(
                    CvImprovementSuggestion.match_result_id == match_id
                )
            ) is None
        if job_description_id is not None:
            assert db.get(JobDescription, job_description_id) is None


def _poll_json(
    client: TestClient,
    path: str,
    headers: dict[str, str],
    *,
    status_field: str,
    attempts: int = 30,
    interval_seconds: float = 1,
) -> dict:
    payload: dict = {}
    for _ in range(attempts):
        response = client.get(path, headers=headers)
        assert response.status_code == 200, response.text
        payload = response.json()
        status = payload.get(status_field)
        if status == "Success":
            return payload
        if status == "Failed":
            pytest.fail(
                f"Live E2E failed at {path}: "
                f"{payload.get('error_message') or 'unknown error'}"
            )
        time.sleep(interval_seconds)
    pytest.fail(f"Live E2E timed out at {path}; last payload: {payload}")
