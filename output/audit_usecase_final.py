from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


PATH = Path(r"C:\Studybase\FitCV_Use_Case_Specification_Final.docx")
doc = Document(PATH)


def iter_blocks(document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


use_cases = {}
current = None
for block in iter_blocks(doc):
    if isinstance(block, Paragraph):
        text = block.text.strip()
        style = block.style.name if block.style else ""
        if style == "Heading 2" and text.startswith("UC-"):
            current = text.split()[0]
            use_cases[current] = {
                "heading": text,
                "paragraphs": [],
                "tables": [],
            }
        elif current and style == "Heading 1":
            current = None
        elif current and text:
            use_cases[current]["paragraphs"].append((style, text))
    elif current:
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in block.rows
        ]
        use_cases[current]["tables"].append(rows)


print(
    f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} "
    f"sections={len(doc.sections)} use_cases={len(use_cases)}"
)

required_markers = [
    "Primary Actor",
    "Precondition",
    "Basic Flow",
    "Alternative",
    "Postcondition",
]
issues = []

for use_case_id, data in use_cases.items():
    all_text_parts = [data["heading"]]
    all_text_parts.extend(text for _, text in data["paragraphs"])
    for table in data["tables"]:
        all_text_parts.extend(cell for row in table for cell in row)
    all_text = "\n".join(all_text_parts)

    marker_status = {
        marker: marker.lower() in all_text.lower()
        for marker in required_markers
    }
    basic_steps = sum(
        1
        for style, text in data["paragraphs"]
        if style == "List Number" or (
            text.split(".", 1)[0].isdigit() and "Step" not in text
        )
    )
    exception_cases = sum(
        1
        for table in data["tables"]
        for row in table
        if row and row[0].startswith("E") and row[0][1:].isdigit()
    )
    nonempty_cells = sum(
        1
        for table in data["tables"]
        for row in table
        for cell in row
        if cell
    )
    missing = [name for name, present in marker_status.items() if not present]
    if missing or basic_steps == 0:
        issues.append((use_case_id, missing, basic_steps))

    print(
        f"{use_case_id}: tables={len(data['tables'])}; "
        f"paragraphs={len(data['paragraphs'])}; cells={nonempty_cells}; "
        f"basic_steps={basic_steps}; exceptions={exception_cases}; "
        f"markers={marker_status}"
    )

print("issues=" + repr(issues))

for use_case_id in ("UC-01", "UC-12"):
    data = use_cases[use_case_id]
    print(f"\n--- {use_case_id} PARAGRAPHS ---")
    for style, text in data["paragraphs"]:
        print(f"[{style}] {text}")
    print(f"--- {use_case_id} TABLES ---")
    for table_index, table in enumerate(data["tables"], start=1):
        print(f"TABLE {table_index}")
        for row in table:
            print(" | ".join(row))
