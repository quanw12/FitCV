from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


path = Path(r"C:\Studybase\FitCV\output\sad_db_work\rup_sad_reference.docx")
doc = Document(path)

print(f"sections={len(doc.sections)} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
for index, section in enumerate(doc.sections, start=1):
    print(
        f"section={index} page={section.page_width.inches:.2f}x{section.page_height.inches:.2f} "
        f"margins={section.top_margin.inches:.2f},{section.right_margin.inches:.2f},"
        f"{section.bottom_margin.inches:.2f},{section.left_margin.inches:.2f} "
        f"header={section.header_distance.inches:.2f} footer={section.footer_distance.inches:.2f}"
    )
    print("  header paragraphs:", [p.text for p in section.header.paragraphs])
    print(
        "  header tables:",
        [[[cell.text for cell in row.cells] for row in table.rows] for table in section.header.tables],
    )
    print("  footer paragraphs:", [p.text for p in section.footer.paragraphs])
    print(
        "  footer tables:",
        [[[cell.text for cell in row.cells] for row in table.rows] for table in section.footer.tables],
    )

print("\nPARAGRAPHS")
for index, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip():
        print(f"P{index:03d}|{paragraph.style.name}|{paragraph.text}")

print("\nTABLES")
for ti, table in enumerate(doc.tables):
    print(f"T{ti:02d}|rows={len(table.rows)}|cols={len(table.columns)}|style={table.style.name if table.style else None}")
    for ri, row in enumerate(table.rows):
        values = [" / ".join(p.text for p in cell.paragraphs) for cell in row.cells]
        print(f"  R{ri:02d}|" + " || ".join(values))

with ZipFile(path) as source:
    names = source.namelist()
    print("\nMEDIA", [name for name in names if name.startswith("word/media/")])
    print("COMMENTS", [name for name in names if "comments" in name])
    root = etree.fromstring(source.read("word/document.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    fields = root.xpath("//w:instrText/text()", namespaces=ns)
    print("FIELDS", fields)
    print("CONTENT_CONTROLS", len(root.xpath("//w:sdt", namespaces=ns)))
