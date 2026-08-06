from hashlib import sha256
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


template = Path(r"C:\Studybase\FitCV\output\rup_restructure_work\rup_template_working.docx")
final = Path(r"C:\Studybase\FitCV_Use_Case_Specification_RUP.docx")
doc = Document(final)

headings_1 = [p.text for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
headings_2 = [p.text for p in doc.paragraphs if p.style and p.style.name == "Heading 2"]
body_text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    body_text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)

assert headings_1 == ["Use-case Model", "Use-case Specifications"], headings_1
assert headings_2 == [
    f"Use-case: UC-{i:02d} - {title}"
    for i, title in enumerate(
        [
            "Upload CV",
            "Analyze CV against Job Description",
            "View AI Improvement Suggestions",
            "Manage CV History",
            "Track Applications",
            "Manage Job Posts (CRUD)",
            "Upload Candidate CV",
            "View Candidate Ranking",
            "Manage Candidate Pipeline",
            "Parse CV",
            "Extract JD Requirements",
            "Calculate Match Score",
        ],
        start=1,
    )
], headings_2

assert len(doc.tables) == 14, len(doc.tables)
spec_tables = doc.tables[2:]
expected_labels = [
    "Use case Name",
    "Brief description",
    "Actors",
    "Basic Flow",
    "Alternative Flows",
    "Pre-conditions",
    "Post-conditions",
]
alternative_count = 0
for index, table in enumerate(spec_tables, start=1):
    labels = [row.cells[0].text.strip() for row in table.rows]
    assert labels == expected_labels, (index, labels)
    assert table.cell(3, 1).text.strip().startswith("1."), index
    alt_text = table.cell(4, 1).text
    assert "Alternative flow 1:" in alt_text, index
    assert "From Step #" in alt_text, index
    assert "Condition:" not in alt_text, index
    assert "System response:" not in alt_text, index
    assert "Return / resume point:" not in alt_text, index
    assert not any(
        paragraph.text.strip().startswith("3.")
        for paragraph in table.cell(4, 1).paragraphs
    ), index
    alternative_count += alt_text.count("Alternative flow ")
    post_text = table.cell(6, 1).text
    assert "Success:" not in post_text, index
    assert "Failure:" not in post_text, index

assert alternative_count == 47, alternative_count

assert "UC-07 Upload Candidate CV <<include>> UC-10 Parse CV" in body_text
assert "Weekly Report Retro Add-on" not in body_text
assert "Review Notes" not in body_text
assert "Common Exception Rules" not in body_text
assert "<Project Name>" not in body_text
assert "[Put an image" not in body_text
assert "<1.0>" not in body_text
assert "<24/07/2026>" not in body_text
assert "\nTrigger\n" not in body_text
assert "\nBusiness Rules\n" not in body_text

assert len(doc.sections) == 2
for section in doc.sections:
    assert round(section.page_width.inches, 2) == 8.50
    assert round(section.page_height.inches, 2) == 11.00
    assert all(
        round(value.inches, 2) == 1.00
        for value in [section.top_margin, section.right_margin, section.bottom_margin, section.left_margin]
    )

with ZipFile(final) as zf:
    document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    settings_xml = zf.read("word/settings.xml").decode("utf-8", errors="ignore")
    assert 'TOC \\o "1-2"' in document_xml
    assert "updateFields" in settings_xml

preserve_parts = [
    "word/styles.xml",
    "word/numbering.xml",
    "word/theme/theme1.xml",
    "word/fontTable.xml",
    "word/webSettings.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
]
with ZipFile(template) as source_zip, ZipFile(final) as final_zip:
    preserve = {}
    for part in preserve_parts:
        source_hash = sha256(source_zip.read(part)).hexdigest()
        final_hash = sha256(final_zip.read(part)).hexdigest()
        preserve[part] = source_hash == final_hash

print("RUP QA PASSED")
print(f"sections={len(doc.sections)} tables={len(doc.tables)} use_cases={len(headings_2)}")
print("preserve_only_part_hashes=" + repr(preserve))
