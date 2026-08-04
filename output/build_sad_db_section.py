from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Studybase\FitCV\output\sad_db_work\rup_sad_reference.docx")
OUTPUT = Path(r"C:\Studybase\FitCV_Software_Architecture_Document_DB_Design.docx")
CORE_ERD = Path(r"C:\Studybase\FitCV\output\sad_db_work\fitcv_erd_core.png")
ANALYSIS_ERD = Path(r"C:\Studybase\FitCV\output\sad_db_work\fitcv_erd_analysis.png")

BLUE = "D9EAF7"
HEADER_BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
WHITE = "FFFFFF"
GRID = "A6A6A6"


ENTITY_ROWS = [
    ("Identity & organization", "industry", "Industry reference data used to classify companies."),
    ("Identity & organization", "company", "Employer profile linked to an industry and to HR accounts and jobs."),
    ("Identity & organization", "account", "Authentication identity, role, profile, company link, and password-reset state."),
    ("Identity & organization", "position", "Normalized job-position reference data."),
    ("Identity & organization", "level", "Normalized seniority-level reference data."),
    ("CV & candidate", "candidate", "Candidate profile for self-service students or HR-created applicants."),
    ("CV & candidate", "cv", "Uploaded CV metadata, ownership, SHA-256 hash, and version history."),
    ("CV & candidate", "cv_parse_result", "Structured output and status produced by the CV parser."),
    ("Recruitment", "job", "Job post, company, owner, position, level, requirements, status, and deadline."),
    ("Recruitment", "job_hr", "Many-to-many assignment of HR accounts to jobs with a role type."),
    ("Recruitment", "job_description", "Reusable or job-linked JD input with raw text and content hash."),
    ("Recruitment", "jd_parse_result", "Structured requirements extracted from a job description."),
    ("Recruitment", "application", "Candidate-to-job application using a selected CV and current hiring stage."),
    ("Student tracking", "tracked_application", "Student-managed external application, status, follow-up, and reminder."),
    ("Student tracking", "tracked_application_note", "Time-ordered notes belonging to a tracked application."),
    ("Student tracking", "tracked_application_status_history", "Audit history for every tracked-application status change."),
    ("AI analysis", "match_result", "CV/JD score, category scores, pass probability, evidence, and algorithm version."),
    ("AI analysis", "cv_improvement_suggestion", "Prioritized AI suggestions derived from one match result."),
    ("AI orchestration", "ai_task", "Asynchronous AI work item with provider, model, status, attempts, and error data."),
]


RELATIONSHIP_ROWS = [
    ("industry 1 — 0..N company", "company.industry_id", "SET NULL", "An industry can classify many companies."),
    ("company 1 — 0..N account", "account.company_id", "SET NULL", "A company can be linked to many internal users."),
    ("account 1 — 0..1 candidate", "candidate.account_id", "SET NULL", "A student account may own one candidate profile."),
    ("account 1 — 0..N candidate", "candidate.created_by_hr_account_id", "SET NULL", "HR can create multiple candidate profiles."),
    ("account/candidate 1 — 0..N cv", "cv.account_id or cv.candidate_id", "CASCADE", "Every CV has at least one owner path."),
    ("cv 1 — 0..N cv_parse_result", "cv_parse_result.cv_id", "CASCADE", "A CV may be parsed repeatedly across parser versions."),
    ("company 1 — 0..N job", "job.company_id", "CASCADE", "Deleting a company removes its jobs."),
    ("account 1 — 0..N job", "job.created_by_account_id", "RESTRICT", "A job retains a valid creator."),
    ("job N — N account", "job_hr(job_id, hr_account_id)", "CASCADE", "HR assignment is resolved by a composite-key bridge."),
    ("account/job 1 — 0..N job_description", "job_description.account_id/job_id", "CASCADE/SET NULL", "A JD belongs to an account and may reference a job."),
    ("job_description 1 — 0..N jd_parse_result", "jd_parse_result.job_description_id", "CASCADE", "A JD may have multiple parse attempts or versions."),
    ("candidate/job/cv 1 — 0..N application", "application candidate/job/cv FKs", "CASCADE/RESTRICT", "The application stores the exact CV used."),
    ("account 1 — 0..N tracked_application", "tracked_application.account_id", "CASCADE", "Student tracking data is account-scoped."),
    ("tracked_application 1 — 0..N note/history", "tracked_application_id", "CASCADE", "Notes and audit history follow the tracked application."),
    ("CV/JD inputs 1 — 0..N match_result", "match_result source FKs", "CASCADE", "A score remains traceable to the inputs and parser outputs."),
    ("application 1 — 0..N match_result", "match_result.application_id", "SET NULL", "A match may support a formal application."),
    ("match_result 1 — 0..N suggestion", "cv_improvement_suggestion.match_result_id", "CASCADE", "Suggestions cannot outlive their source analysis."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if value and keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.08):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(anchor, text="", style=None, bold_lead=None):
    paragraph = anchor.insert_paragraph_before(style=style)
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    set_paragraph_spacing(paragraph)
    return paragraph


def add_bullet(anchor, text):
    paragraph = add_text(anchor, f"• {text}", "Normal")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.15)
    return paragraph


def add_heading(anchor, text):
    paragraph = add_text(anchor, text, "Heading 3")
    set_keep_with_next(paragraph)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def format_table(table, widths=None):
    table.style = "Normal Table"
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), GRID)
        borders.append(node)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=0, line=1.0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
        if row_index == 0:
            set_repeat_table_header(row)
            for cell in row.cells:
                set_cell_shading(cell, HEADER_BLUE)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        elif row_index % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "EDF4FB")
    return table


def add_table_before(doc, anchor, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    format_table(table, widths)
    anchor._p.addprevious(table._tbl)
    spacer = anchor.insert_paragraph_before()
    set_paragraph_spacing(spacer, after=2)
    return table


def add_picture(anchor, path, caption):
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(path), width=Inches(6.35))
    caption_paragraph = anchor.insert_paragraph_before(style="Normal")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    set_paragraph_spacing(caption_paragraph, after=6)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not CORE_ERD.exists() or not ANALYSIS_ERD.exists():
        raise FileNotFoundError("ER diagrams must be generated before building the document.")

    doc = Document(SOURCE)
    component_heading = next(p for p in doc.paragraphs if p.text.strip() == "Component: xyz")
    next_component = next(p for p in doc.paragraphs if p.text.strip() == "Deployment")

    component_heading.text = "Component: Database"
    component_heading.style = "Heading 2"

    paragraphs = doc.paragraphs
    start = next(index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "Component: Database")
    stop = next(index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "Deployment")
    for paragraph in list(paragraphs[start + 1:stop]):
        remove_paragraph(paragraph)

    add_heading(next_component, "Responsibilities")
    add_text(
        next_component,
        "The Database component provides durable, relational storage for FitCV. It preserves user and organization data, CV versions, job and job-description data, applications, AI parsing outputs, match evidence, improvement suggestions, and asynchronous AI task state. The component is accessed by the FastAPI backend through SQLAlchemy repositories; the React frontend does not connect to the database directly.",
    )
    add_bullet(next_component, "Maintain authentication identities and role-aware links to companies and candidates.")
    add_bullet(next_component, "Preserve uploaded CV metadata and immutable version history while keeping the binary file in external file storage.")
    add_bullet(next_component, "Store recruitment jobs, HR ownership, candidate applications, student application tracking, reminders, notes, and status history.")
    add_bullet(next_component, "Persist AI input/output provenance so every match score and suggestion can be traced to the exact CV, JD, parser output, and algorithm version.")

    add_heading(next_component, "Technology and Physical Design")
    add_text(
        next_component,
        "The physical schema targets MySQL 8.0 or later. All tables use the InnoDB engine for transactions and foreign-key enforcement, and utf8mb4 with utf8mb4_unicode_ci for multilingual CV and JD text. Primary identifiers use BIGINT UNSIGNED AUTO_INCREMENT unless a bridge table requires a composite key. Timestamps use DATETIME values generated by the database where appropriate.",
    )
    add_bullet(next_component, "Application access path: React/Vite → HTTPS REST API → FastAPI service → repository → SQLAlchemy → MySQL.")
    add_bullet(next_component, "Large CV files are stored outside MySQL; the cv table stores the path, MIME type, size, SHA-256 hash, and version metadata.")
    add_bullet(next_component, "Structured AI payloads and evidence use JSON columns so the schema can retain explainable model output without losing relational ownership and lifecycle rules.")

    add_heading(next_component, "Entity-Relationship Models")
    add_text(
        next_component,
        "The ER design is split into two readable views of one physical schema. Solid lines represent enforced foreign keys. The dashed ai_task link is logical because ai_task uses the generic pair task_type and resource_id rather than a physical foreign key.",
    )
    page_break = next_component.insert_paragraph_before()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    add_picture(
        next_component,
        CORE_ERD,
        "Figure 1. FitCV ER model A — identity, recruitment, CV ownership, and the hiring pipeline.",
    )
    add_picture(
        next_component,
        ANALYSIS_ERD,
        "Figure 2. FitCV ER model B — parsing, matching, AI suggestions, and student application tracking.",
    )

    add_heading(next_component, "Entity Catalog")
    add_text(
        next_component,
        "Table 1 summarizes every physical table in database/full_schema.sql. The grouping is architectural only; all entities belong to the same FitCV MySQL schema.",
    )
    add_table_before(
        doc,
        next_component,
        ("Domain", "Entity", "Primary responsibility"),
        ENTITY_ROWS,
        (1.35, 1.75, 3.25),
    )

    add_heading(next_component, "Key Relationships and Lifecycle Rules")
    add_text(
        next_component,
        "Table 2 records the principal cardinalities and deletion behavior. CASCADE is used for true dependent records, SET NULL preserves optional historical context, and RESTRICT prevents deletion when a referenced record must remain valid.",
    )
    add_table_before(
        doc,
        next_component,
        ("Relationship", "Foreign key", "On delete", "Design meaning"),
        RELATIONSHIP_ROWS,
        (1.45, 1.55, 0.75, 2.60),
    )

    add_heading(next_component, "Integrity Constraints")
    add_bullet(next_component, "A CV must belong to an account or a candidate: account_id IS NOT NULL OR candidate_id IS NOT NULL.")
    add_bullet(next_component, "CV version_number is unique per account, allowing reliable history and latest-version retrieval.")
    add_bullet(next_component, "Every match_result must reference either a job or a standalone job_description.")
    add_bullet(next_component, "Overall, skill, experience, education, soft-skill, and pass-probability values are constrained to the inclusive range 0–100.")
    add_bullet(next_component, "The tuple (cv_parse_id, jd_parse_id, algorithm_version) is unique, preventing duplicate results for the same parsed inputs and scoring algorithm.")
    add_bullet(next_component, "Reference names and identities are protected by unique keys, including industry_name, position.abbreviation, level_name, and account.email.")

    add_heading(next_component, "Indexes and Query Support")
    add_text(
        next_component,
        "Foreign-key and workflow indexes support the dominant FitCV access patterns: accounts by company or role; candidate ownership; latest CV version; parser output by CV or JD; jobs by company or creator; applications by candidate or job; tracked applications by account, date, status, or reminder; match history by CV and generation time; and ordered suggestions by match and type. ai_task is indexed by task_type, resource_id, and created_at for background-worker polling and audit lookup.",
    )

    add_heading(next_component, "Security and Data Handling")
    add_bullet(next_component, "Passwords and password-reset codes are stored only as hashes; reset expiry is stored separately.")
    add_bullet(next_component, "Database credentials are supplied through environment variables and are never committed to source control.")
    add_bullet(next_component, "Backend authorization restricts data by authenticated account and role before repository queries are executed.")
    add_bullet(next_component, "SHA-256 values support duplicate detection and traceability for CV files and JD text; they are not a replacement for access control.")
    add_bullet(next_component, "AI scores are decision-support information, not guaranteed hiring outcomes; evidence_json and algorithm_version support review and reproducibility.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
