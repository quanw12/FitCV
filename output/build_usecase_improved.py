from pathlib import Path
import math

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Studybase")
OUT = ROOT / "FitCV_Use_Case_Specification_Final.docx"
DIAGRAM = Path(r"C:\Studybase\FitCV\output\fitcv_overall_usecase_diagram.png")


BLUE = RGBColor(37, 99, 235)
SLATE = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
LIGHT = "EEF6FF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if row_i == 0:
                set_cell_shading(cell, "DBEAFE")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = SLATE


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=SLATE)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
    style_table(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def create_restarted_numbering(doc, style_name="List Number"):
    numbering = doc.part.numbering_part.element
    style_id = doc.styles[style_name].style_id
    abstract_num_id = None
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        styles = abstract_num.findall(".//" + qn("w:pStyle"))
        if any(node.get(qn("w:val")) == style_id for node in styles):
            abstract_num_id = abstract_num.get(qn("w:abstractNumId"))
            break
    if abstract_num_id is None:
        raise ValueError(f"No numbering definition found for {style_name}.")

    existing_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_numbered(doc, items):
    num_id = create_restarted_numbering(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(num_id))
        num_pr.append(level)
        num_pr.append(number)
        p_pr.append(num_pr)
        p.add_run(item)


def validate_basic_flow_numbering(doc):
    flow_num_ids = {}
    current_uc = None
    in_basic_flow = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 2" and paragraph.text.startswith("UC-"):
            current_uc = paragraph.text.split(" - ", 1)[0]
            in_basic_flow = False
        elif paragraph.style.name == "Heading 3":
            in_basic_flow = paragraph.text == "Basic Flow"
        elif in_basic_flow and paragraph.style.name == "List Number":
            num_id = paragraph._p.pPr.numPr.numId.val
            flow_num_ids.setdefault(current_uc, set()).add(int(num_id))

    expected = {f"UC-{index:02d}" for index in range(1, 13)}
    if set(flow_num_ids) != expected:
        raise ValueError(f"Basic Flow numbering is missing use cases: {expected - set(flow_num_ids)}")
    if any(len(ids) != 1 for ids in flow_num_ids.values()):
        raise ValueError(f"Each Basic Flow must use exactly one numbering sequence: {flow_num_ids}")
    if len({next(iter(ids)) for ids in flow_num_ids.values()}) != len(expected):
        raise ValueError("Each use case must have its own restarted numbering sequence.")


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = BLUE
    r2 = p.add_run(body)
    r2.font.name = "Arial"
    r2.font.size = Pt(9)
    style_table(table)
    doc.add_paragraph()


def draw_actor(draw, x, y, label, font):
    draw.ellipse((x - 14, y, x + 14, y + 28), outline=(15, 23, 42), width=2)
    draw.line((x, y + 28, x, y + 80), fill=(15, 23, 42), width=2)
    draw.line((x - 32, y + 48, x + 32, y + 48), fill=(15, 23, 42), width=2)
    draw.line((x, y + 80, x - 28, y + 122), fill=(15, 23, 42), width=2)
    draw.line((x, y + 80, x + 28, y + 122), fill=(15, 23, 42), width=2)
    bbox = draw.textbbox((0, 0), label, font=font)
    draw.text((x - (bbox[2] - bbox[0]) / 2, y + 132), label, fill=(15, 23, 42), font=font)


def ellipse(draw, box, text, font, fill=(255, 255, 255), outline=(37, 99, 235)):
    draw.ellipse(box, fill=fill, outline=outline, width=3)
    lines = text.split("\n")
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + (len(lines) - 1) * 4
    cy = (box[1] + box[3]) / 2 - total_h / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text(((box[0] + box[2]) / 2 - (bbox[2] - bbox[0]) / 2, cy), line, fill=(15, 23, 42), font=font)
        cy += (bbox[3] - bbox[1]) + 4


def line(draw, p1, p2, dashed=False, color=(148, 163, 184)):
    if not dashed:
        draw.line((*p1, *p2), fill=color, width=2)
        return
    x1, y1 = p1
    x2, y2 = p2
    segments = 24
    for i in range(segments):
        if i % 2 == 0:
            xa = x1 + (x2 - x1) * i / segments
            ya = y1 + (y2 - y1) * i / segments
            xb = x1 + (x2 - x1) * (i + 1) / segments
            yb = y1 + (y2 - y1) * (i + 1) / segments
            draw.line((xa, ya, xb, yb), fill=color, width=2)


def arrowhead(draw, p1, p2, color=(71, 85, 105)):
    x1, y1 = p1
    x2, y2 = p2
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 16
    spread = 0.48
    a1 = angle + math.pi - spread
    a2 = angle + math.pi + spread
    points = [
        (x2, y2),
        (x2 + size * math.cos(a1), y2 + size * math.sin(a1)),
        (x2 + size * math.cos(a2), y2 + size * math.sin(a2)),
    ]
    draw.line((points[0], points[1]), fill=color, width=2)
    draw.line((points[0], points[2]), fill=color, width=2)


def polyline(draw, points, color=(148, 163, 184), width=3):
    draw.line(points, fill=color, width=width, joint="curve")


def use_case_ellipse(draw, box, uc_id, label, id_font, label_font, fill, outline):
    draw.ellipse(box, fill=fill, outline=outline, width=4)
    center_x = (box[0] + box[2]) / 2
    center_y = (box[1] + box[3]) / 2
    id_box = draw.textbbox((0, 0), uc_id, font=id_font)
    label_box = draw.multiline_textbbox((0, 0), label, font=label_font, spacing=4, align="center")
    total_height = (id_box[3] - id_box[1]) + 8 + (label_box[3] - label_box[1])
    id_y = center_y - total_height / 2
    draw.text(
        (center_x - (id_box[2] - id_box[0]) / 2, id_y),
        uc_id,
        fill=(37, 99, 235),
        font=id_font,
    )
    label_y = id_y + (id_box[3] - id_box[1]) + 8
    draw.multiline_text(
        (center_x, label_y),
        label,
        fill=(15, 23, 42),
        font=label_font,
        spacing=4,
        anchor="ma",
        align="center",
    )


def create_diagram():
    width, height = 3000, 1750
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        label_font = ImageFont.truetype("arial.ttf", 38)
        id_font = ImageFont.truetype("arialbd.ttf", 30)
        small = ImageFont.truetype("arial.ttf", 29)
        title = ImageFont.truetype("arialbd.ttf", 50)
        group_font = ImageFont.truetype("arialbd.ttf", 32)
    except OSError:
        label_font = id_font = small = title = group_font = ImageFont.load_default()

    blue = (37, 99, 235)
    slate = (15, 23, 42)
    muted = (100, 116, 139)
    association = (148, 163, 184)
    relation = (71, 85, 105)
    user_fill = (255, 255, 255)
    support_fill = (248, 250, 252)

    boundary = (410, 70, 2590, 1660)
    draw.rounded_rectangle(boundary, radius=30, outline=blue, width=5)
    draw.text((465, 110), "FitCV System Boundary", fill=blue, font=title)

    student_panel = (500, 225, 1510, 1070)
    support_panel = (1630, 225, 2495, 1070)
    hr_panel = (500, 1130, 2495, 1575)
    for box in (student_panel, support_panel, hr_panel):
        draw.rounded_rectangle(box, radius=24, fill=(250, 252, 255), outline=(219, 229, 242), width=3)
    draw.text((545, 255), "STUDENT / JOB SEEKER", fill=blue, font=group_font)
    draw.text((1675, 255), "SHARED PROCESSING", fill=blue, font=group_font)
    draw.text((545, 1160), "HR / RECRUITMENT", fill=blue, font=group_font)

    actors = {
        "Student": (165, 555),
        "HR / Recruiter": (165, 1280),
        "File Storage": (2825, 305),
        "AI Matching\nService": (2825, 650),
        "Hiring Manager": (2825, 1280),
    }
    for label, (x, y) in actors.items():
        draw_actor(draw, x, y, label, small)

    ucs = {
        "UC-05": ((650, 330, 1190, 455), "Track Applications", user_fill, blue),
        "UC-01": ((650, 480, 1190, 605), "Upload CV", user_fill, blue),
        "UC-04": ((650, 630, 1190, 755), "Manage CV History", user_fill, blue),
        "UC-02": ((650, 760, 1190, 895), "Analyze CV against JD", user_fill, blue),
        "UC-03": ((650, 950, 1190, 1060), "View AI Suggestions", user_fill, blue),
        "UC-10": ((1770, 450, 2310, 585), "Parse CV", support_fill, muted),
        "UC-11": ((1770, 690, 2310, 825), "Extract JD Requirements", support_fill, muted),
        "UC-12": ((1770, 865, 2310, 1000), "Calculate Match Score", support_fill, muted),
        "UC-06": ((570, 1270, 1015, 1410), "Manage Job Posts\n(CRUD)", user_fill, blue),
        "UC-07": ((1080, 1270, 1510, 1410), "Upload Candidate\nCV", user_fill, blue),
        "UC-08": ((1580, 1270, 2015, 1410), "View Candidate\nRanking", user_fill, blue),
        "UC-09": ((2070, 1270, 2490, 1410), "Manage Candidate\nPipeline", user_fill, blue),
    }

    # Plain UML actor associations. Orthogonal routing keeps them away from labels.
    student_targets = [(650, 392), (650, 542), (650, 692), (650, 847), (650, 997)]
    student_starts = [(225, 585), (225, 605), (225, 625), (225, 645), (225, 665)]
    for start, target in zip(student_starts, student_targets):
        polyline(draw, [start, target], association, 2)

    polyline(draw, [(225, 1325), (570, 1340)], association, 2)
    polyline(draw, [(225, 1350), (1080, 1305)], association, 2)
    polyline(draw, [(225, 1375), (1580, 1290)], association, 2)

    # Supporting actors use short routes along the right margin.
    polyline(draw, [(2765, 375), (2545, 375), (2545, 517), (2310, 517)], association, 3)
    polyline(draw, [(2765, 720), (2460, 720), (2310, 757)], association, 3)
    polyline(draw, [(2765, 720), (2460, 720), (2310, 932)], association, 3)
    polyline(draw, [(2765, 1335), (2490, 1340)], association, 2)
    polyline(draw, [(2765, 1365), (2015, 1290)], association, 2)

    # Draw associations first, then ovals so no line crosses use-case text.
    for uc_id, (box, label, fill, outline) in ucs.items():
        use_case_ellipse(draw, box, uc_id, label, id_font, label_font, fill, outline)

    # Include relationships point from the base use case to mandatory reused behavior.
    line(draw, (1190, 542), (1770, 517), dashed=True, color=relation)
    arrowhead(draw, (1190, 542), (1770, 517), color=relation)
    draw.text((1350, 495), "<<include>>", fill=relation, font=small)

    # Candidate CV uploads reuse the same mandatory parsing use case.
    candidate_parse_route = [(1295, 1270), (1295, 1105), (1570, 1105), (1570, 610), (1770, 570)]
    for start, end in zip(candidate_parse_route, candidate_parse_route[1:]):
        line(draw, start, end, dashed=True, color=relation)
    arrowhead(draw, candidate_parse_route[-2], candidate_parse_route[-1], color=relation)
    draw.text((1320, 1060), "<<include>>", fill=relation, font=small)

    line(draw, (1190, 802), (1770, 757), dashed=True, color=relation)
    arrowhead(draw, (1190, 802), (1770, 757), color=relation)
    draw.text((1350, 740), "<<include>>", fill=relation, font=small)

    line(draw, (1190, 853), (1770, 932), dashed=True, color=relation)
    arrowhead(draw, (1190, 853), (1770, 932), color=relation)
    draw.text((1350, 895), "<<include>>", fill=relation, font=small)

    # The optional extension points back to its base use case.
    line(draw, (1125, 950), (1125, 895), dashed=True, color=relation)
    arrowhead(draw, (1125, 950), (1125, 895), color=relation)
    draw.text((1165, 905), "<<extend>>", fill=relation, font=small)

    draw.text(
        (540, 1600),
        "Legend: actor association = solid line   |   <<include>> / <<extend>> = dashed arrow",
        fill=muted,
        font=small,
    )
    DIAGRAM.parent.mkdir(parents=True, exist_ok=True)
    img.save(DIAGRAM)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLUE if name == "Heading 1" else SLATE
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)


def make_section_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def make_section_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def use_case(doc, uc_id, name, actor, goal, trigger, pre, post_success, post_failure, basic, exceptions, rules):
    doc.add_heading(f"{uc_id} - {name}", level=2)
    add_table(
        doc,
        ["Field", "Description"],
        [
            ["Use Case ID", uc_id],
            ["Use Case Name", name],
            ["Primary Actor", actor],
            ["Goal", goal],
            ["Trigger", trigger],
        ],
        [1.5, 5.5],
    )
    doc.add_heading("Preconditions", level=3)
    add_bullets(doc, pre)
    doc.add_heading("Postconditions", level=3)
    add_table(doc, ["Result", "Postcondition"], [["Success", post_success], ["Failure", post_failure]], [1.4, 5.6])
    doc.add_heading("Basic Flow", level=3)
    add_numbered(doc, basic)
    doc.add_heading("Alternative / Exception Flows", level=3)
    add_table(doc, ["Case", "Condition", "System Response", "Return / Resume Point"], exceptions, [0.75, 1.7, 2.5, 2.0])
    doc.add_heading("Business Rules", level=3)
    add_bullets(doc, rules)


def build():
    create_diagram()
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    make_section_portrait(section)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("FitCV Use Case Specification")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.color.rgb = BLUE
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Improved PA2 Draft - Standard Use Case Model and Specification")
    sr.font.name = "Arial"
    sr.font.size = Pt(11)
    sr.font.color.rgb = MUTED

    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Project", "FitCV - AI-powered CV Screening and Job Readiness Platform"],
            ["Document Type", "Use Case Specification"],
            ["Version", "Improved draft based on instructor feedback"],
        ["Main Actors", "Student / Job Seeker, HR / Recruiter, Hiring Manager, AI Matching Service, File Storage"],
        ],
        [1.7, 5.5],
    )

    doc.add_heading("1. Document Purpose", level=1)
    doc.add_paragraph(
        "This document defines the main use cases of the FitCV platform. It explains the actors, system boundary, "
        "use-case relationships, detailed flows, exception flows, and business rules needed for PA2."
    )
    add_note(
        doc,
        "Feedback addressed",
        "This version adds one overall use-case diagram, uses UML-style relationship labels, adds HR CRUD management, and makes exception flows return to clear steps.",
    )

    doc.add_heading("2. System Overview", level=1)
    doc.add_paragraph(
        "FitCV supports students who want to improve their CV before applying for jobs and HR users who need to manage job posts, "
        "candidate CVs, rankings, and pipeline decisions. The AI feature provides recommendations only; it does not make final hiring decisions."
    )

    doc.add_heading("3. Actors", level=1)
    add_table(
        doc,
        ["Actor", "Description"],
        [
            ["Student / Job Seeker", "Uploads CVs, compares CVs with job descriptions, reviews suggestions, manages CV history, and tracks job applications."],
            ["HR / Recruiter", "Creates, views, updates, and closes/deletes job posts; uploads candidate CVs; reviews rankings; and manages candidate pipeline stages."],
            ["Hiring Manager", "Reviews shortlisted candidates, match summaries, and pipeline recommendations."],
            ["AI Matching Service", "External/supporting service that extracts CV/JD information and produces match scores, summaries, and suggestions."],
            ["File Storage", "Stores uploaded CV files and returns file references to the system."],
        ],
        [1.7, 5.5],
    )

    diagram_section = doc.add_section(WD_SECTION.NEW_PAGE)
    make_section_landscape(diagram_section)
    doc.add_heading("4. Use Case Model", level=1)
    doc.add_paragraph(
        "The following model shows the common FitCV system boundary. Actor associations are plain UML lines. "
        "Mandatory reused behavior is shown with <<include>>, while optional behavior is shown with <<extend>>."
    )
    doc.add_picture(str(DIAGRAM), width=Inches(9.8))
    cap = doc.add_paragraph("Figure 1. Overall Use-case Diagram for FitCV")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(
        doc,
        ["Relationship", "Meaning in FitCV"],
        [
            ["UC-01 Upload CV <<include>> UC-10 Parse CV", "Parsing is mandatory after a valid CV is uploaded."],
            ["UC-07 Upload Candidate CV <<include>> UC-10 Parse CV", "Parsing is mandatory after a candidate CV is uploaded by HR."],
            ["UC-02 Analyze CV against JD <<include>> UC-11 Extract JD Requirements", "The system must extract structured job requirements before scoring."],
            ["UC-02 Analyze CV against JD <<include>> UC-12 Calculate Match Score", "A completed analysis always calculates and stores a match score."],
            ["UC-03 View AI Suggestions <<extend>> UC-02 Analyze CV against JD", "Viewing suggestions is optional behavior available after a completed analysis."],
        ],
        [2.4, 4.8],
    )

    portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
    make_section_portrait(portrait_section)
    doc.add_heading("5. Use Case List", level=1)
    add_table(
        doc,
        ["ID", "Use Case", "Primary Actor", "Priority"],
        [
            ["UC-01", "Upload CV", "Student / Job Seeker", "High"],
            ["UC-02", "Analyze CV against Job Description", "Student / Job Seeker", "Highest"],
            ["UC-03", "View AI Improvement Suggestions", "Student / Job Seeker", "High"],
            ["UC-04", "Manage CV History", "Student / Job Seeker", "Medium"],
            ["UC-05", "Track Applications", "Student / Job Seeker", "High"],
            ["UC-06", "Manage Job Posts (CRUD)", "HR / Recruiter", "Highest"],
            ["UC-07", "Upload Candidate CV", "HR / Recruiter", "High"],
            ["UC-08", "View Candidate Ranking", "HR / Recruiter", "Highest"],
            ["UC-09", "Manage Candidate Pipeline", "HR / Recruiter / Hiring Manager", "High"],
            ["UC-10", "Parse CV", "Student / Job Seeker or HR / Recruiter", "High"],
            ["UC-11", "Extract JD Requirements", "Student / Job Seeker", "High"],
            ["UC-12", "Calculate Match Score", "Student / Job Seeker", "Highest"],
        ],
        [0.8, 3.2, 2.0, 1.0],
    )

    doc.add_heading("6. Detailed Use Case Specifications", level=1)
    use_case(
        doc,
        "UC-01",
        "Upload CV",
        "Student / Job Seeker",
        "The student uploads a CV file so it can be stored and parsed for matching.",
        "The student clicks Upload CV.",
        ["The student is logged in.", "The CV file is available on the student's device.", "The file format should be PDF or DOCX."],
        "The CV file is stored, a CV record is created, and parsing is started or completed.",
        "No CV record is saved, or the uploaded CV is marked as Failed with a clear reason.",
        ["Student opens the CV upload area.", "Student selects a CV file.", "System validates file type and file size.", "System uploads the file to storage.", "System creates a CV record linked to the student.", "System sends the CV to the parser.", "System shows upload and parse status."],
        [
            ["E1", "File is not PDF or DOCX.", "System rejects the file and displays the accepted formats.", "Return to Step 2: choose another file."],
            ["E2", "File is larger than 10MB.", "System displays the size limit and does not upload the file.", "Return to Step 2: choose a smaller file."],
            ["E3", "Storage upload fails.", "System shows a retry option and keeps the selected file in the UI if possible.", "Return to Step 4: retry upload."],
            ["E4", "CV parser fails.", "System keeps the CV record but marks parse status as Failed.", "Return to Step 2 or Step 6: upload another file or retry parsing."],
        ],
        ["Only PDF and DOCX files are accepted.", "Maximum file size is 10MB.", "A student may upload multiple CV versions.", "Only one CV should be marked as latest at a time."],
    )

    use_case(
        doc,
        "UC-02",
        "Analyze CV against Job Description",
        "Student / Job Seeker",
        "The student compares a CV with a job description to receive a match score, evidence, and improvement suggestions.",
        "The student clicks Analyze Match.",
        ["The student is logged in.", "A valid CV is selected or uploaded.", "A job description is provided by text input, file upload, or selected job post."],
        "The match result is saved and displayed with overall score, score breakdown, strengths, weaknesses, and suggestions.",
        "No completed match result is saved; the system explains what must be corrected before analysis can continue.",
        ["Student opens the CV & JD Match Analyzer page.", "Student selects or uploads a CV.", "Student enters, uploads, or selects a job description.", "System validates that both CV and JD are available.", "System retrieves parsed CV content.", "System extracts JD requirements.", "System sends CV and JD data to the AI Matching Service.", "System saves and displays the match result."],
        [
            ["E1", "JD is empty or too short.", "System asks the student to provide a complete job description.", "Return to Step 3: edit JD input."],
            ["E2", "CV parsing is still processing.", "System shows loading status and disables analysis temporarily.", "Resume at Step 5 when parsing completes."],
            ["E3", "AI service timeout or quota error.", "System shows retry message and does not save a fake completed result.", "Return to Step 7: retry analysis."],
            ["E4", "AI output is incomplete or invalid.", "System marks the match as Failed and stores diagnostic status.", "Return to Step 7 after service recovers or model output is corrected."],
        ],
        ["Score ranges from 0 to 100.", "80-100 is Strong Match, 50-79 is Moderate Match, and 0-49 is Weak Match.", "The system must show both number and text label.", "AI output is a recommendation, not a guaranteed hiring decision."],
    )

    use_case(
        doc,
        "UC-03",
        "View AI Improvement Suggestions",
        "Student / Job Seeker",
        "The student reviews categorized suggestions to improve the CV for a target job.",
        "The student opens a completed match result or clicks View Suggestions.",
        ["The student is logged in.", "At least one completed match result exists.", "The result contains suggestion data."],
        "The student can view prioritized suggestions with explanation and suggested wording.",
        "The system shows why suggestions are unavailable and offers the next possible action.",
        ["Student opens a saved match result.", "System loads match score and suggestion data.", "System groups suggestions by Skill, Experience, Education, Keyword, or Format.", "Student expands a suggestion.", "System displays explanation, evidence, and suggested revision."],
        [
            ["E1", "No suggestions exist.", "System shows an empty state and explains that no improvement data was generated.", "Return to Step 1: select another result or rerun analysis."],
            ["E2", "The result does not belong to the student.", "System denies access.", "End use case and return to authorized results list."],
            ["E3", "The CV or JD changed after analysis.", "System warns the user that suggestions may be outdated.", "Return to UC-02 Step 7: rerun analysis if needed."],
        ],
        ["Suggestions must be linked to a match result.", "Suggestions should include Low, Medium, or High priority.", "The system should not automatically rewrite or submit the CV without user action."],
    )

    use_case(
        doc,
        "UC-04",
        "Manage CV History",
        "Student / Job Seeker",
        "The student reviews uploaded CV versions, selects a version for analysis, and removes an obsolete version when allowed.",
        "The student opens CV History.",
        ["The student is logged in.", "The student owns zero or more CV records."],
        "The CV history reflects the student's latest permitted action and keeps version information consistent.",
        "No unauthorized or invalid change is applied; the existing CV history remains unchanged.",
        ["Student opens CV History.", "System loads only CV records owned by the student.", "System displays file name, version number, upload date, parse status, and latest-version indicator.", "Student selects a CV to view details or use in analysis.", "Student may choose Delete for an obsolete CV.", "System requests confirmation and checks dependent match/application records.", "System applies the permitted action.", "System refreshes the ordered CV history."],
        [
            ["E1", "The student has no uploaded CV.", "System shows an empty state with an Upload CV action.", "Return to UC-01 Step 2."],
            ["E2", "The selected CV is still processing.", "System shows its current parse status and disables analysis for that version.", "Resume at Step 4 when parsing succeeds."],
            ["E3", "The selected CV does not belong to the student.", "System denies access without exposing CV data.", "End the use case and return to Step 2."],
            ["E4", "The CV has dependent records that prevent hard deletion.", "System explains the dependency and offers archive/keep behavior if supported.", "Return to Step 5: choose another CV or cancel."],
        ],
        ["CV history must be owner-scoped.", "Version numbers must be unique within a student's CV history.", "Only one CV version should be marked as latest.", "Deleting a CV must not leave broken match or application references."],
    )

    use_case(
        doc,
        "UC-05",
        "Track Applications",
        "Student / Job Seeker",
        "The student records job applications, updates their status, adds follow-up notes, and monitors reminders.",
        "The student opens Application Tracker or clicks Add Application.",
        ["The student is logged in.", "The student has permission to manage only their own application records."],
        "The application, note, reminder, and status history data reflect the student's submitted changes.",
        "No invalid or unauthorized change is stored; the tracker shows a clear recovery action.",
        ["Student opens Application Tracker.", "System loads the student's application list and status summary.", "Student creates or selects an application.", "Student enters or updates company, position, application date, source, status, and optional follow-up date.", "System validates required fields and ownership.", "System saves the application.", "If the status changed, system records the previous and new status in history.", "Student may add, edit, or delete a note.", "System refreshes filters, counts, and due-reminder indicators."],
        [
            ["E1", "Required company, position, or date data is missing.", "System highlights invalid fields and prevents saving.", "Return to Step 4: correct the application form."],
            ["E2", "The requested status is not supported.", "System rejects the update and displays valid statuses.", "Return to Step 4: select Applied, Screening, Interview, Offer, or Rejected."],
            ["E3", "The application or note belongs to another account.", "System denies access without exposing the record.", "End the action and return to Step 2."],
            ["E4", "Saving a note, status, or reminder fails.", "System keeps the entered data in the form and offers Retry.", "Return to Step 6 or Step 8: retry the failed save."],
        ],
        ["Valid application statuses are Applied, Screening, Interview, Offer, and Rejected.", "Every status change must create a history entry.", "Reminder dates and notes are optional.", "Offer and Rejected applications should not create stale follow-up alerts.", "All tracker records must be owner-scoped."],
    )

    use_case(
        doc,
        "UC-06",
        "Manage Job Posts (CRUD)",
        "HR / Recruiter",
        "The HR user manages job posts through create, view, update, and close/delete operations.",
        "The HR user opens Job Post Management.",
        ["The HR user is logged in.", "The HR user has permission for the company workspace.", "Company information exists."],
        "The selected CRUD action is completed and the job list reflects the latest state.",
        "The action is not applied and the system shows validation or permission errors.",
        ["HR opens the Job Post Management page.", "System displays the job list with status, title, deadline, and actions.", "HR chooses Create, View, Edit, Close, or Delete.", "System opens the correct form or confirmation dialog.", "HR submits the action.", "System validates required fields and permissions.", "System creates, updates, closes, deletes, or displays the job post.", "System returns HR to the refreshed job list."],
        [
            ["E1", "Required fields are missing when creating or updating.", "System highlights missing fields and prevents saving.", "Return to Step 4: correct the form."],
            ["E2", "HR has no permission for the company/job.", "System blocks the action and shows permission error.", "Return to Step 2: show only authorized jobs."],
            ["E3", "Delete/close is selected for a job with active candidates.", "System asks for confirmation or suggests closing instead of hard delete.", "Return to Step 4: confirm or cancel."],
            ["E4", "Database save fails.", "System shows retry message and keeps form data.", "Return to Step 5: retry submit."],
        ],
        ["Published jobs require title, description or requirements, company, creator, and status.", "Job status should include Draft, Published, Closed, and optionally Archived.", "CRUD actions must be permission-checked.", "Deleting should be restricted when historical application records depend on the job."],
    )

    use_case(
        doc,
        "UC-07",
        "Upload Candidate CV",
        "HR / Recruiter",
        "The HR user uploads a candidate CV for screening even if the candidate does not have a FitCV account.",
        "The HR user clicks Upload Candidate CV.",
        ["The HR user is logged in.", "The HR user has permission to manage candidates for a selected job.", "A job post exists or is selected."],
        "Candidate, CV, and application records are created or updated.",
        "The candidate CV is not attached to the job and the system explains the reason.",
        ["HR opens the candidate management page for a job.", "HR clicks Upload Candidate CV.", "HR enters candidate information if available.", "HR selects a CV file.", "System validates file type and size.", "System creates or updates candidate data.", "System uploads the CV and creates an application record.", "System starts CV parsing."],
        [
            ["E1", "Candidate email already exists.", "System suggests linking to the existing candidate.", "Return to Step 3: confirm candidate identity."],
            ["E2", "No job is selected.", "System asks HR to select a job.", "Return to Step 1: choose a job."],
            ["E3", "Invalid CV file.", "System rejects the file and explains the accepted format/size.", "Return to Step 4: choose another file."],
            ["E4", "CV parsing fails.", "System keeps the application but marks analysis status as Failed.", "Return to Step 8: retry parsing or upload a cleaner CV."],
        ],
        ["HR-uploaded candidates are not required to have an account.", "A candidate may apply to multiple jobs.", "Each application must reference the CV used for that application."],
    )

    use_case(
        doc,
        "UC-08",
        "View Candidate Ranking",
        "HR / Recruiter",
        "The HR user views candidates ranked by AI match score for a selected job.",
        "The HR user opens Candidate Ranking for a job.",
        ["The HR user is logged in.", "The HR user has permission to view the job.", "The job has at least one candidate application."],
        "The system displays ranked candidates with score, label, stage, strengths, weaknesses, and actions.",
        "Ranking cannot be displayed and the system shows the recovery action.",
        ["HR selects a job post.", "System loads candidate applications.", "System retrieves or requests match results.", "System sorts candidates by overall match score.", "System displays ranking with score breakdown and stage.", "HR opens a candidate detail view.", "HR updates candidate stage if needed."],
        [
            ["E1", "No candidates exist.", "System shows an empty state with Upload Candidate CV action.", "Return to UC-07 Step 2."],
            ["E2", "A candidate has no match result.", "System offers to run analysis for that candidate.", "Return to Step 3 after analysis completes."],
            ["E3", "AI analysis failed.", "System keeps the candidate visible and marks analysis as Failed.", "Return to Step 3: retry analysis."],
            ["E4", "HR lacks permission.", "System blocks access.", "End use case and return to authorized job list."],
        ],
        ["Ranking supports HR decisions but does not automatically accept or reject candidates.", "The interface must show both score and label.", "Candidate stage may include Applied, Screening, Interview, Offer, Hired, or Rejected."],
    )

    use_case(
        doc,
        "UC-09",
        "Manage Candidate Pipeline",
        "HR / Recruiter / Hiring Manager",
        "Authorized recruitment users move candidates through hiring stages, record review context, and preserve stage history.",
        "The user opens the candidate pipeline for a selected job.",
        ["The user is logged in.", "The user has permission to view or update the selected job.", "At least one candidate application exists."],
        "The permitted candidate stage change is saved, displayed, and recorded in pipeline history.",
        "The candidate remains in the previous stage and the system explains why the requested action was not applied.",
        ["User selects a job and opens its pipeline.", "System groups candidates by current stage.", "User opens a candidate card or drags it to a permitted stage.", "System asks for confirmation when the transition is sensitive.", "System validates role, job ownership, and transition rules.", "System saves the new stage and audit information.", "System refreshes pipeline counts and candidate detail.", "Hiring Manager or HR reviews the updated shortlist."],
        [
            ["E1", "The selected job has no candidates.", "System shows an empty pipeline with an Upload Candidate CV action.", "Return to UC-07 Step 2."],
            ["E2", "The user lacks permission to change the candidate stage.", "System leaves the candidate unchanged and shows a permission error.", "Return to Step 2 in read-only mode."],
            ["E3", "The requested stage transition is invalid.", "System explains the allowed next stages.", "Return to Step 3: choose a valid stage."],
            ["E4", "Concurrent update or database save fails.", "System reloads the latest candidate stage and asks the user to retry.", "Return to Step 3 with refreshed data."],
        ],
        ["Pipeline changes must be permission-checked and auditable.", "AI ranking is decision support only; a human user approves every stage change.", "A candidate's stage history must preserve old stage, new stage, actor, and timestamp.", "Sensitive transitions such as Offer, Hired, or Rejected should require confirmation."],
    )

    use_case(
        doc,
        "UC-10",
        "Parse CV",
        "Student / Job Seeker or HR / Recruiter (supporting: AI Matching Service and File Storage)",
        "The system converts a stored CV into validated text and structured evidence that later use cases can reuse.",
        "UC-01 or UC-07 includes this use case after a valid CV upload.",
        ["A valid PDF or DOCX file has been stored.", "A CV record exists with Pending or Processing status."],
        "Parsed text/evidence is linked to the CV record and parse status becomes Success.",
        "The CV is marked Failed with a non-sensitive diagnostic message and can be retried.",
        ["System loads the stored CV file.", "System verifies the file can be read.", "System extracts text directly from PDF/DOCX or requests OCR when configured and necessary.", "System normalizes the extracted text.", "System extracts structured CV evidence.", "System validates that evidence is grounded in the extracted text.", "System saves the parse result.", "System updates the CV parse status to Success."],
        [
            ["E1", "The stored file is missing or corrupted.", "System marks parsing as Failed and tells the user to upload the file again.", "Return to UC-01 Step 2 or UC-07 Step 4, depending on the calling flow."],
            ["E2", "A scanned PDF has no text layer and OCR is unavailable.", "System explains that OCR is unavailable and requests a text-based PDF or DOCX.", "Return to UC-01 Step 2 or UC-07 Step 4, depending on the calling flow."],
            ["E3", "OCR or AI extraction times out.", "System marks the task Failed without inventing extracted content.", "Return to Step 3: retry when the service is available."],
            ["E4", "Extracted evidence fails validation.", "System discards the invalid structured result and records Failed status.", "Return to Step 5: retry extraction."],
        ],
        ["The original file and parse result must remain linked.", "Only evidence present in the source CV may be saved.", "Failure must never be reported as a successful parse.", "CV content must be protected according to access and privacy rules."],
    )

    use_case(
        doc,
        "UC-11",
        "Extract JD Requirements",
        "Student / Job Seeker (supporting: AI Matching Service)",
        "The system converts a complete job description into structured skills, experience, education, and soft-skill requirements.",
        "UC-02 includes this use case before calculating a match score.",
        ["A job description has been entered, uploaded, or selected.", "The text meets minimum completeness validation."],
        "Validated structured JD requirements are available to the scoring use case.",
        "No requirements are accepted and UC-02 remains incomplete with a corrective message.",
        ["System receives the JD text.", "System validates that the JD is complete enough to analyze.", "System normalizes the text.", "System requests structured requirement extraction.", "System validates categories and quoted evidence against the JD.", "System removes unsupported or duplicate items.", "System returns the validated requirements to UC-02."],
        [
            ["E1", "The JD is empty or too short.", "System requests a more complete job description.", "Return to UC-02 Step 3."],
            ["E2", "The extraction service is unavailable or rate-limited.", "System shows Failed/Retry status and does not create placeholder requirements.", "Return to Step 4: retry later."],
            ["E3", "The service returns invalid structured output.", "System rejects the output and records a validation error.", "Return to Step 4 after correcting the request or service output."],
            ["E4", "Extracted evidence is not present in the JD.", "System removes unsupported items; if no valid requirements remain, extraction fails.", "Return to Step 4 or UC-02 Step 3."],
        ],
        ["Every extracted requirement must be grounded in the JD text.", "Duplicate requirements should be normalized.", "Missing categories must remain empty rather than being invented.", "The stored JD must remain linked to its parse result."],
    )

    use_case(
        doc,
        "UC-12",
        "Calculate Match Score",
        "Student / Job Seeker (supporting: AI Matching Service)",
        "The system compares validated CV evidence with validated JD requirements and calculates an explainable match result.",
        "UC-02 includes this use case after UC-10 and UC-11 provide valid inputs.",
        ["The selected CV has a successful parse result.", "Validated JD requirements are available.", "The user owns the selected CV."],
        "A match result is stored with overall score, category breakdown, evidence, gaps, and a text label.",
        "No successful score is stored; UC-02 shows Failed status and a retry path.",
        ["System loads validated CV evidence and JD requirements.", "System compares evidence by Skills, Experience, Education, and Soft Skills.", "System calculates category scores.", "System applies configured category weights and redistributes missing-category weight when required.", "System calculates an overall score from 0 to 100.", "System assigns Strong, Moderate, or Weak Match label.", "System saves score breakdown, evidence, and gaps.", "System returns the match result to UC-02."],
        [
            ["E1", "The CV parse result is missing or failed.", "System blocks scoring and requests CV parsing.", "Return to UC-10 Step 3."],
            ["E2", "The JD has no validated requirements.", "System blocks scoring and requests JD extraction.", "Return to UC-11 Step 2."],
            ["E3", "The comparison service times out or returns invalid data.", "System records Failed status and does not expose a fabricated score.", "Return to Step 2: retry comparison."],
            ["E4", "The computed total is outside 0-100 or category data is inconsistent.", "System rejects the result and records a validation error.", "Return to Step 3: recalculate after correction."],
        ],
        ["Default weights are Skills 45%, Experience 30%, Education 15%, and Soft Skills 10%.", "If a JD category is absent, its weight is redistributed across available categories.", "80-100 is Strong Match, 50-79 is Moderate Match, and 0-49 is Weak Match.", "The result is explainable decision support and must not automatically accept or reject a candidate."],
    )

    doc.add_heading("7. Common Exception Rules", level=1)
    add_table(
        doc,
        ["Exception Type", "Standard Handling", "Return Point"],
        [
            ["Authentication required", "Redirect user to login or show authentication error.", "Return to the original page after successful login."],
            ["Permission denied", "Block access and show only resources owned by or assigned to the user.", "Return to the authorized list page."],
            ["Validation error", "Highlight invalid fields and explain the correction.", "Return to the exact form step where the invalid data was entered."],
            ["AI service error", "Do not store misleading successful results; show retry or failed status.", "Return to the AI request step after the service is available."],
            ["No data", "Show an empty state and next recommended action.", "Return to the create/upload action for that module."],
        ],
        [1.5, 3.5, 2.2],
    )

    doc.add_heading("8. Weekly Report Retro Add-on", level=1)
    doc.add_paragraph("Use the following block in the weekly report to satisfy the instructor's retro requirement.")
    add_table(
        doc,
        ["Went Well", "Could Be Improved"],
        [
            ["The team completed the main FitCV student and HR use cases for the MVP scope.", "The use-case model should be unified into one diagram instead of being split into separate partial diagrams."],
            ["The AI matching flow was connected to practical student needs such as CV/JD comparison and improvement suggestions.", "Exception flows should always state where the user returns after an error."],
            ["The HR workflow includes candidate upload and ranking, which supports the recruitment process.", "HR CRUD management should be written explicitly, not only implied through Create Job Post."],
        ],
        [3.5, 3.5],
    )

    doc.add_heading("9. Review Notes", level=1)
    doc.add_paragraph(
        "Compared with the earlier draft, this improved version fixes the missing overall use-case model, clarifies UML relationship notation, "
        "adds HR CRUD management, and strengthens exception flows with clear return points."
    )

    validate_basic_flow_numbering(doc)
    doc.save(OUT)
    print(OUT)
    print(DIAGRAM)


if __name__ == "__main__":
    build()
