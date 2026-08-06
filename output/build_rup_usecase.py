from copy import deepcopy
from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Studybase")
WORK = ROOT / "FitCV" / "output" / "rup_restructure_work"
TEMPLATE = WORK / "rup_template_working.docx"
SOURCE = ROOT / "FitCV_Use_Case_Specification_Final.docx"
DIAGRAM = ROOT / "FitCV" / "output" / "fitcv_overall_usecase_diagram.png"
OUTPUT = ROOT / "FitCV_Use_Case_Specification_RUP.docx"

PROJECT = "FitCV"
TEAM = "FitCV Team"
VERSION = "1.0"
DOC_ID = "FITCV-PA2-UCS-1.0"
DOC_DATE = "24/07/2026"
REVISION_DATE = "24/Jul/2026"

ALTERNATIVE_FROM_STEPS = {
    "UC-01": [3, 3, 4, 6],
    "UC-02": [4, 5, 7, 7],
    "UC-03": [2, 2, 2],
    "UC-04": [2, 4, 2, 6],
    "UC-05": [5, 5, 5, 6],
    "UC-06": [6, 6, 6, 7],
    "UC-07": [6, 1, 5, 8],
    "UC-08": [2, 3, 3, 1],
    "UC-09": [2, 5, 5, 6],
    "UC-10": [2, 3, 3, 6],
    "UC-11": [2, 4, 5, 5],
    "UC-12": [1, 1, 2, 5],
}


def iter_blocks(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def table_map(table):
    result = {}
    for row in table.rows:
        if len(row.cells) >= 2:
            result[row.cells[0].text.strip()] = row.cells[1].text.strip()
    return result


def extract_use_cases(path):
    doc = Document(path)
    cases = []
    current = None
    subsection = None

    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            style = block.style.name if block.style else ""
            if style == "Heading 1" and current:
                cases.append(current)
                current = None
                subsection = None
                continue
            if style == "Heading 2" and text.startswith("UC-"):
                if current:
                    cases.append(current)
                uc_id, title = text.split(" - ", 1)
                current = {
                    "id": uc_id,
                    "title": title,
                    "preconditions": [],
                    "basic_flow": [],
                    "rules": [],
                    "exceptions": [],
                }
                subsection = None
                continue
            if not current:
                continue
            if style == "Heading 3":
                subsection = text
                continue
            if not text:
                continue
            if subsection == "Preconditions":
                current["preconditions"].append(text)
            elif subsection == "Basic Flow":
                current["basic_flow"].append(text)
            elif subsection == "Business Rules":
                current["rules"].append(text)
        elif current:
            first = block.cell(0, 0).text.strip() if block.rows else ""
            if first == "Field":
                metadata = table_map(block)
                current["actor"] = metadata["Primary Actor"]
                current["goal"] = metadata["Goal"]
                current["trigger"] = metadata["Trigger"]
            elif first == "Result":
                outcomes = table_map(block)
                current["success"] = outcomes["Success"]
                current["failure"] = outcomes["Failure"]
            elif first == "Case":
                for row in block.rows[1:]:
                    current["exceptions"].append(
                        {
                            "case": row.cells[0].text.strip(),
                            "condition": row.cells[1].text.strip(),
                            "response": row.cells[2].text.strip(),
                            "return": row.cells[3].text.strip(),
                        }
                    )

    if current:
        cases.append(current)

    if [c["id"] for c in cases] != [f"UC-{i:02d}" for i in range(1, 13)]:
        raise ValueError("Expected UC-01 through UC-12 in order.")
    for case in cases:
        required = ("actor", "goal", "trigger", "success", "failure")
        if any(not case.get(key) for key in required):
            raise ValueError(f"Incomplete metadata in {case['id']}.")
        if not case["basic_flow"] or not case["exceptions"]:
            raise ValueError(f"Incomplete flow in {case['id']}.")
    return cases


def remove_all_body_content(doc):
    body = doc.element.body
    final_sect_pr = body.sectPr
    first_break = None
    for child in list(body):
        if child.tag == qn("w:p") and child.find("./w:pPr/w:sectPr", namespaces=child.nsmap) is not None:
            first_break = deepcopy(child)
            break
    if first_break is None:
        raise ValueError("The RUP template section break was not found.")
    for child in list(body):
        if child is not final_sect_pr:
            body.remove(child)
    return first_break, final_sect_pr


def add_paragraph_before_final_sectpr(doc, text="", style=None):
    return doc.add_paragraph(text, style=style)


def insert_preserved_section_break(doc, break_paragraph, final_sect_pr):
    body = doc.element.body
    body.insert(body.index(final_sect_pr), break_paragraph)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    return p


def add_label(cell, text):
    p = clear_cell(cell)
    run = p.add_run(text)
    run.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_plain(cell, text):
    p = clear_cell(cell)
    p.add_run(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_lines(cell, lines, numbered=False, bullet=False):
    clear_cell(cell)
    first = True
    for index, line in enumerate(lines, start=1):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(2)
        if numbered:
            p.add_run(f"{index}.  {line}")
        elif bullet:
            p.add_run(f"•  {line}")
        else:
            p.add_run(line)


def lower_first(text):
    if not text:
        return text
    return text[0].lower() + text[1:]


def narrative_continue(text):
    text = re.sub(r"\bStep (\d+)", r"Step #\1", text.strip())
    if text.startswith("Return to "):
        target = text[len("Return to ") :]
        if ":" in target:
            location, action = target.split(":", 1)
            location = location.strip()
            action = action.strip().rstrip(".")
            if location.startswith("Step #"):
                return f"Continue at {location} of the Basic Flow to {action}."
            return f"Continue at {location} to {action}."
        if target.startswith("Step #"):
            return f"Continue at {target.rstrip('.')} of the Basic Flow."
        return f"Continue at {target}"
    if text.startswith("Resume at "):
        target = text[len("Resume at ") :].rstrip(".")
        if target.startswith("Step #"):
            return f"Continue at {target} of the Basic Flow."
        return f"Continue at {target}."
    return text


def add_exception_flows(cell, case_id, exceptions):
    clear_cell(cell)
    from_steps = ALTERNATIVE_FROM_STEPS[case_id]
    if len(from_steps) != len(exceptions):
        raise ValueError(f"Alternative-flow branch-step mapping mismatch for {case_id}.")
    for index, (exc, from_step) in enumerate(zip(exceptions, from_steps), start=1):
        if index > 1:
            spacer = cell.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
        heading = cell.paragraphs[0] if index == 1 else cell.add_paragraph()
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run(f"Alternative flow {index}: {exc['condition']}")
        run.bold = True
        narrative_steps = [
            (
                f"From Step #{from_step} of the Basic Flow, "
                f"{lower_first(exc['condition'].rstrip('.'))}. {exc['response']}"
            ),
            narrative_continue(exc["return"]),
        ]
        for step, value in enumerate(narrative_steps, start=1):
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)
            p.add_run(f"{step}.  {value}")


def add_spec_table(doc, case):
    rows = [
        "Use case Name",
        "Brief description",
        "Actors",
        "Basic Flow",
        "Alternative Flows",
        "Pre-conditions",
        "Post-conditions",
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row, label in zip(table.rows, rows):
        set_cell_width(row.cells[0], 1.45)
        set_cell_width(row.cells[1], 5.05)
        add_label(row.cells[0], label)
        prevent_row_split(row)

    add_plain(table.cell(0, 1), f"{case['id']} - {case['title']}")
    add_plain(table.cell(1, 1), case["goal"])
    add_plain(table.cell(2, 1), case["actor"])
    add_lines(table.cell(3, 1), case["basic_flow"], numbered=True)
    add_exception_flows(table.cell(4, 1), case["id"], case["exceptions"])
    add_lines(table.cell(5, 1), case["preconditions"], bullet=True)
    add_plain(table.cell(6, 1), case["success"])

    return table


def add_revision_table(doc):
    table = doc.add_table(rows=5, cols=4)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.35, 0.85, 3.55, 1.25]
    headers = ["Date", "Version", "Description", "Author"]
    for i, (header, width) in enumerate(zip(headers, widths)):
        set_cell_width(table.cell(0, i), width)
        p = clear_cell(table.cell(0, i))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(header).bold = True
    values = [REVISION_DATE, VERSION, "Restructured into the standard RUP Use-Case Specification template.", TEAM]
    for i, value in enumerate(values):
        set_cell_width(table.cell(1, i), widths[i])
        add_plain(table.cell(1, i), value)
    set_repeat_table_header(table.rows[0])
    return table


def add_field(paragraph, instruction, cached_text=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = cached_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def enable_field_updates(doc):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def replace_header_footer(doc):
    sections = doc.sections
    if len(sections) != 2:
        raise ValueError(f"Expected two sections, got {len(sections)}.")

    cover_header = sections[0].header
    for p in cover_header.paragraphs:
        if "<Team Name>" in p.text:
            p.text = TEAM
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    header = sections[1].header
    if not header.tables:
        raise ValueError("Template metadata header table is missing.")
    table = header.tables[0]
    table.cell(0, 0).text = PROJECT
    table.cell(0, 1).text = f"Version:          {VERSION}"
    table.cell(1, 0).text = "Use-Case Specification"
    table.cell(1, 1).text = f"Date:  {DOC_DATE}"
    table.cell(2, 0).text = DOC_ID

    footer = sections[1].footer
    if not footer.tables:
        raise ValueError("Template footer table is missing.")
    footer_table = footer.tables[0]
    footer_table.cell(0, 0).text = "Confidential"
    footer_table.cell(0, 1).text = f"© {TEAM}, 2026"
    # Preserve the original PAGE field in the third cell.


def add_relationship_table(doc):
    relationships = [
        ("UC-01 Upload CV <<include>> UC-10 Parse CV", "Parsing is mandatory after a valid CV is uploaded."),
        ("UC-07 Upload Candidate CV <<include>> UC-10 Parse CV", "Parsing is mandatory after a candidate CV is uploaded by HR."),
        ("UC-02 Analyze CV against JD <<include>> UC-11 Extract JD Requirements", "JD extraction is mandatory before scoring."),
        ("UC-02 Analyze CV against JD <<include>> UC-12 Calculate Match Score", "A completed analysis calculates and stores a match score."),
        ("UC-03 View AI Suggestions <<extend>> UC-02 Analyze CV against JD", "Suggestions are optional after a completed analysis."),
    ]
    table = doc.add_table(rows=1 + len(relationships), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Relationship", "Meaning in FitCV"]
    for index, header in enumerate(headers):
        add_label(table.cell(0, index), header)
    set_repeat_table_header(table.rows[0])
    for row, values in zip(table.rows[1:], relationships):
        add_plain(row.cells[0], values[0])
        add_plain(row.cells[1], values[1])
        prevent_row_split(row)
    return table


def build():
    cases = extract_use_cases(SOURCE)
    doc = Document(TEMPLATE)
    doc.core_properties.title = "FitCV Use-Case Specification"
    doc.core_properties.subject = "PA2 - RUP Use-Case Model and Specifications"
    doc.core_properties.author = TEAM
    doc.core_properties.last_modified_by = TEAM
    doc.core_properties.comments = "Restructured from the verified FitCV use-case content using the supplied RUP template."

    section_break, final_sect_pr = remove_all_body_content(doc)

    # Cover page.
    doc.add_paragraph(PROJECT, style="Title")
    doc.add_paragraph("Use-Case Specification", style="Title")
    doc.add_paragraph("", style="Normal")
    version_p = doc.add_paragraph(style="Title")
    version_p.add_run(f"Version {VERSION}")
    insert_preserved_section_break(doc, section_break, final_sect_pr)

    # Revision history.
    doc.add_paragraph("Revision History", style="Title")
    add_revision_table(doc)
    doc.add_page_break()

    # Table of contents.
    doc.add_paragraph("Table of Contents", style="Title")
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "Table of Contents - updates automatically when opened.")
    doc.add_page_break()

    # Use-case model.
    doc.add_heading("Use-case Model", level=1)
    doc.add_paragraph(
        "The following model shows the complete FitCV system boundary. Actor associations use solid UML lines. "
        "Mandatory reused behavior uses <<include>>, while optional behavior uses <<extend>>."
    )
    picture_p = doc.add_paragraph()
    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_p.add_run().add_picture(str(DIAGRAM), width=Inches(6.35))
    caption = doc.add_paragraph("Figure 1. Overall Use-case Diagram for FitCV")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_relationship_table(doc)
    doc.add_page_break()

    # Detailed specifications.
    doc.add_heading("Use-case Specifications", level=1)
    for index, case in enumerate(cases, start=1):
        heading = doc.add_heading(f"Use-case: {case['id']} - {case['title']}", level=2)
        if index > 1:
            heading.paragraph_format.page_break_before = True
        add_spec_table(doc, case)

    enable_field_updates(doc)
    replace_header_footer(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
